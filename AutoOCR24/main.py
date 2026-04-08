import html
import os
import shutil
import statistics
import tempfile
import threading
import time
from dataclasses import dataclass
from pathlib import Path
from tkinter import filedialog, messagebox, ttk
import tkinter as tk

import pdfplumber
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.edge.options import Options
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select, WebDriverWait


APP_NAME = "AutoOCR24"
PDF24_OCR_URL = "https://tools.pdf24.org/es/ocr-pdf"
PDF24_LANGUAGES = ("cat", "spa")
ICON_PATH = Path(r"C:\Users\solso\Documents\New project\assets\logo.ico")
EDGE_CANDIDATES = [
    Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
    Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
]

COLORS = {
    "red": "#C8102E",
    "red_dark": "#9E0E24",
    "cream": "#F6F0E8",
    "paper": "#FFFDF9",
    "line": "#E5D8CB",
    "charcoal": "#1F1F1F",
    "muted": "#675F58",
}


@dataclass
class OCRPage:
    number: int
    text: str


def find_edge_binary() -> str:
    for candidate in EDGE_CANDIDATES:
        if candidate.exists():
            return str(candidate)
    raise FileNotFoundError("No he trobat Microsoft Edge instal·lat en este ordinador.")


def structured_page_text(page) -> str:
    words = page.extract_words(x_tolerance=2, y_tolerance=3, keep_blank_chars=False)
    if not words:
        return (page.extract_text() or "").strip()

    words = sorted(words, key=lambda item: (round(item["top"], 1), item["x0"]))
    lines = []
    for word in words:
        top = word["top"]
        bottom = word["bottom"]
        height = bottom - top
        if not lines:
            lines.append({"top": top, "bottom": bottom, "words": [word]})
            continue
        previous = lines[-1]
        tolerance = max(3.0, min(height, previous["bottom"] - previous["top"]) * 0.55)
        if abs(top - previous["top"]) <= tolerance:
            previous["words"].append(word)
            previous["bottom"] = max(previous["bottom"], bottom)
        else:
            lines.append({"top": top, "bottom": bottom, "words": [word]})

    text_lines = []
    heights = [line["bottom"] - line["top"] for line in lines]
    median_height = statistics.median(heights) if heights else 12
    previous_bottom = None

    for line in lines:
        ordered_words = sorted(line["words"], key=lambda item: item["x0"])
        average_char_widths = [
            (item["x1"] - item["x0"]) / max(len(item["text"]), 1)
            for item in ordered_words
            if item["text"]
        ]
        average_char_width = statistics.mean(average_char_widths) if average_char_widths else 6

        chunks = []
        previous_word = None
        for word in ordered_words:
            if previous_word is not None:
                gap = word["x0"] - previous_word["x1"]
                if gap > average_char_width * 5:
                    chunks.append("    ")
                else:
                    chunks.append(" ")
            chunks.append(word["text"])
            previous_word = word

        line_text = "".join(chunks).strip()
        if previous_bottom is not None:
            gap = line["top"] - previous_bottom
            if gap > median_height * 0.9:
                text_lines.append("")
        text_lines.append(line_text)
        previous_bottom = line["bottom"]

    return "\n".join(text_lines).strip()


def extract_document_pages(pdf_path: Path) -> list[OCRPage]:
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = structured_page_text(page)
            pages.append(OCRPage(number=page_number, text=text))
    return pages


def dismiss_consent_overlay(driver, timeout=20):
    try:
        WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, "iframe#consentManagerOverlay"))
        )
    except Exception:
        return

    try:
        driver.switch_to.default_content()
        iframe = driver.find_element(By.CSS_SELECTOR, "iframe#consentManagerOverlay")
        driver.switch_to.frame(iframe)
        button = WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((By.CSS_SELECTOR, "a.acceptAll"))
        )
        driver.execute_script("arguments[0].click();", button)
    finally:
        driver.switch_to.default_content()
    try:
        WebDriverWait(driver, timeout).until(
            EC.invisibility_of_element_located((By.CSS_SELECTOR, "iframe#consentManagerOverlay"))
        )
    except Exception:
        pass


def wait_for_download(directory: Path, timeout=120) -> Path:
    started = time.time()
    last_size = {}
    stable_counts = {}
    while time.time() - started < timeout:
        pdfs = [item for item in directory.glob("*.pdf") if item.is_file()]
        for pdf_file in pdfs:
            size = pdf_file.stat().st_size
            if last_size.get(pdf_file) == size:
                stable_counts[pdf_file] = stable_counts.get(pdf_file, 0) + 1
            else:
                stable_counts[pdf_file] = 0
            last_size[pdf_file] = size
            if stable_counts[pdf_file] >= 2 and size > 0:
                return pdf_file
        time.sleep(1)
    raise TimeoutError("No s'ha pogut descarregar el PDF OCR en el temps esperat.")


class PDF24OCRClient:
    def __init__(self, progress_callback):
        self.progress_callback = progress_callback

    def progress(self, text):
        if self.progress_callback:
            self.progress_callback(text)

    def ocr_pdf(self, input_pdf: Path) -> Path:
        edge_binary = find_edge_binary()
        working_dir = Path(tempfile.mkdtemp(prefix="autooocr24_"))
        download_dir = working_dir / "downloads"
        download_dir.mkdir(exist_ok=True)

        options = Options()
        options.use_chromium = True
        options.binary_location = edge_binary
        options.add_argument("--headless=new")
        options.add_argument("--disable-gpu")
        options.add_argument("--window-size=1600,1600")
        options.add_experimental_option(
            "prefs",
            {
                "download.default_directory": str(download_dir),
                "download.prompt_for_download": False,
                "download.directory_upgrade": True,
                "plugins.always_open_pdf_externally": True,
            },
        )

        driver = webdriver.Edge(options=options)
        wait = WebDriverWait(driver, 120)
        try:
            self.progress("Obrint PDF24 Tools...")
            driver.get(PDF24_OCR_URL)
            dismiss_consent_overlay(driver)

            self.progress("Pujant PDF...")
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "input[type=file]")))
            driver.find_element(By.CSS_SELECTOR, "input[type=file]").send_keys(str(input_pdf.resolve()))
            time.sleep(2)

            self.progress("Configurant idiomes OCR: català + castellà...")
            for value in PDF24_LANGUAGES:
                checkbox = driver.find_element(By.CSS_SELECTOR, f'input[type=checkbox][value="{value}"]')
                if not checkbox.is_selected():
                    driver.execute_script("arguments[0].click();", checkbox)

            Select(driver.find_element(By.CSS_SELECTOR, 'select[name="outputType"]')).select_by_value("pdf")

            self.progress("Iniciant OCR en segon pla...")
            driver.execute_script("arguments[0].click();", driver.find_element(By.CSS_SELECTOR, "button.ocrPdf"))
            driver.execute_script("arguments[0].click();", driver.find_element(By.CSS_SELECTOR, "button.submitBtn"))
            wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "iframe.workerFrame")))

            download_clicked = False
            for _ in range(80):
                time.sleep(4)
                driver.switch_to.default_content()
                dismiss_consent_overlay(driver, timeout=3)
                frame = driver.find_element(By.CSS_SELECTOR, "iframe.workerFrame")
                driver.switch_to.frame(frame)

                result_view = driver.find_elements(By.CSS_SELECTOR, "#resultView")
                if result_view and result_view[0].is_displayed():
                    if not download_clicked:
                        self.progress("OCR acabat. Descarregant PDF OCR...")
                        download_button = driver.find_element(By.CSS_SELECTOR, "#downloadTool")
                        driver.execute_script("arguments[0].click();", download_button)
                        download_clicked = True
                else:
                    status = ""
                    status_nodes = driver.find_elements(By.CSS_SELECTOR, "#processingView .status")
                    if status_nodes and status_nodes[0].is_displayed():
                        status = status_nodes[0].text.strip()
                    if status:
                        self.progress(status)

                pdf_files = [item for item in download_dir.glob("*.pdf") if item.is_file()]
                if pdf_files:
                    break

            downloaded_pdf = wait_for_download(download_dir)
            final_path = working_dir / f"{input_pdf.stem}_ocr.pdf"
            shutil.copy2(downloaded_pdf, final_path)
            self.progress("PDF OCR descarregat correctament.")
            return final_path
        finally:
            driver.quit()


def export_docx(pages: list[OCRPage], output_path: Path):
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    for page_index, page in enumerate(pages):
        blocks = [block.strip() for block in page.text.split("\n\n")]
        for block in blocks:
            paragraph = document.add_paragraph()
            lines = [line.rstrip() for line in block.splitlines() if line.strip() or len(blocks) == 1]
            for index, line in enumerate(lines):
                run = paragraph.add_run(line)
                if index < len(lines) - 1:
                    run.add_break(WD_BREAK.LINE)
        if page_index < len(pages) - 1:
            document.add_page_break()

    document.save(output_path)


def export_clean_pdf(pages: list[OCRPage], output_path: Path):
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "CleanText",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        spaceAfter=8,
    )
    story = []
    for page_index, page in enumerate(pages):
        blocks = [block.strip() for block in page.text.split("\n\n") if block.strip()]
        for block in blocks:
            escaped = "<br/>".join(html.escape(line) for line in block.splitlines())
            story.append(Paragraph(escaped, normal))
            story.append(Spacer(1, 6))
        if page_index < len(pages) - 1:
            story.append(PageBreak())

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=52,
        rightMargin=52,
        topMargin=52,
        bottomMargin=52,
    )
    document.build(story)


class AutoOCR24App:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_NAME)
        self.root.geometry("1500x920")
        self.root.minsize(1320, 760)
        self.root.configure(bg=COLORS["cream"])

        self.source_pdf: Path | None = None
        self.ocr_pdf: Path | None = None
        self.pages: list[OCRPage] = []
        self.current_page_index: int | None = None
        self.dirty = False
        self.suspend_dirty = False
        self.worker_thread = None

        self.pdf_path_var = tk.StringVar()
        self.ocr_path_var = tk.StringVar(value="Encara no s'ha generat.")
        self.status_var = tk.StringVar(value="A punt.")

        self.configure_styles()
        self.build_ui()

    def configure_styles(self):
        style = ttk.Style(self.root)
        if "vista" in style.theme_names():
            style.theme_use("vista")

        style.configure("Root.TFrame", background=COLORS["cream"])
        style.configure("Card.TFrame", background=COLORS["paper"])
        style.configure("Header.TFrame", background=COLORS["red"])
        style.configure("HeaderTitle.TLabel", background=COLORS["red"], foreground="white", font=("Segoe UI Semibold", 22))
        style.configure("HeaderSub.TLabel", background=COLORS["red"], foreground="white", font=("Segoe UI", 10))
        style.configure("Section.TLabel", background=COLORS["paper"], foreground=COLORS["charcoal"], font=("Segoe UI Semibold", 10))
        style.configure("Body.TLabel", background=COLORS["paper"], foreground=COLORS["charcoal"], font=("Segoe UI", 10))
        style.configure("Status.TLabel", background=COLORS["cream"], foreground=COLORS["charcoal"], font=("Segoe UI Semibold", 10))
        style.configure("Accent.TButton", font=("Segoe UI Semibold", 10))
        style.configure("Treeview", rowheight=28, font=("Segoe UI", 10), fieldbackground="white", background="white", foreground=COLORS["charcoal"])
        style.configure("Treeview.Heading", font=("Segoe UI Semibold", 10))
        style.map("Treeview", background=[("selected", COLORS["red"])], foreground=[("selected", "white")])

    def build_ui(self):
        header = ttk.Frame(self.root, style="Header.TFrame", padding=(20, 18))
        header.pack(fill="x")
        ttk.Label(header, text=APP_NAME, style="HeaderTitle.TLabel").pack(anchor="w")
        ttk.Label(
            header,
            text="OCR amb PDF24 Tools en segon pla, revisió per pàgines i exportació a DOCX/PDF net.",
            style="HeaderSub.TLabel",
        ).pack(anchor="w", pady=(4, 0))

        top = ttk.Frame(self.root, style="Root.TFrame", padding=(16, 14, 16, 8))
        top.pack(fill="x")
        top_card = ttk.Frame(top, style="Card.TFrame", padding=14)
        top_card.pack(fill="x")

        ttk.Label(top_card, text="PDF origen", style="Section.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Entry(top_card, textvariable=self.pdf_path_var, width=92).grid(row=0, column=1, sticky="ew", padx=8, pady=4)
        ttk.Button(top_card, text="Buscar", command=self.pick_pdf, style="Accent.TButton").grid(row=0, column=2, padx=4)
        ttk.Button(top_card, text="Processar OCR", command=self.process_ocr, style="Accent.TButton").grid(row=0, column=3, padx=4)

        ttk.Label(top_card, text="PDF OCR temporal", style="Section.TLabel").grid(row=1, column=0, sticky="w")
        ttk.Entry(top_card, textvariable=self.ocr_path_var, width=92, state="readonly").grid(row=1, column=1, sticky="ew", padx=8, pady=4)
        ttk.Button(top_card, text="Obrir PDF OCR", command=self.open_ocr_pdf).grid(row=1, column=2, padx=4)
        ttk.Button(top_card, text="Reextraure text", command=self.reload_text_from_ocr).grid(row=1, column=3, padx=4)
        top_card.columnconfigure(1, weight=1)

        controls = ttk.Frame(self.root, style="Root.TFrame", padding=(16, 0, 16, 10))
        controls.pack(fill="x")
        controls_card = ttk.Frame(controls, style="Card.TFrame", padding=10)
        controls_card.pack(fill="x")
        ttk.Button(controls_card, text="Aplicar canvis de pàgina", command=self.apply_current_page, style="Accent.TButton").pack(side="left", padx=4)
        ttk.Button(controls_card, text="Exportar a DOCX", command=self.export_to_docx, style="Accent.TButton").pack(side="right", padx=4)
        ttk.Button(controls_card, text="Exportar a PDF net", command=self.export_to_pdf, style="Accent.TButton").pack(side="right", padx=4)

        body = ttk.PanedWindow(self.root, orient="horizontal")
        body.pack(fill="both", expand=True, padx=16, pady=(0, 10))

        left = ttk.Frame(body, style="Card.TFrame", padding=10)
        center = ttk.Frame(body, style="Card.TFrame", padding=10)
        right = ttk.Frame(body, style="Card.TFrame", padding=10)
        body.add(left, weight=2)
        body.add(center, weight=5)
        body.add(right, weight=2)

        ttk.Label(left, text="Pàgines detectades", style="Section.TLabel").pack(anchor="w", pady=(0, 8))
        self.page_tree = ttk.Treeview(left, columns=("page", "preview"), show="headings", height=24)
        self.page_tree.heading("page", text="Pàgina")
        self.page_tree.heading("preview", text="Vista prèvia")
        self.page_tree.column("page", width=80, anchor="w")
        self.page_tree.column("preview", width=240, anchor="w")
        self.page_tree.pack(side="left", fill="both", expand=True)
        self.page_tree.bind("<<TreeviewSelect>>", self.on_page_select)
        page_scroll = ttk.Scrollbar(left, orient="vertical", command=self.page_tree.yview)
        page_scroll.pack(side="right", fill="y")
        self.page_tree.configure(yscrollcommand=page_scroll.set)

        ttk.Label(center, text="Editor de text", style="Section.TLabel").pack(anchor="w", pady=(0, 8))
        self.editor = tk.Text(
            center,
            wrap="word",
            undo=True,
            font=("Consolas", 11),
            bg="white",
            fg=COLORS["charcoal"],
            insertbackground=COLORS["red"],
            relief="flat",
            padx=12,
            pady=12,
        )
        self.editor.pack(side="left", fill="both", expand=True)
        self.editor.bind("<<Modified>>", self.on_editor_modified)
        editor_scroll = ttk.Scrollbar(center, orient="vertical", command=self.editor.yview)
        editor_scroll.pack(side="right", fill="y")
        self.editor.configure(yscrollcommand=editor_scroll.set)

        ttk.Label(right, text="Estat del procés", style="Section.TLabel").pack(anchor="w", pady=(0, 8))
        self.log_box = tk.Text(
            right,
            wrap="word",
            font=("Segoe UI", 10),
            bg="white",
            fg=COLORS["charcoal"],
            relief="flat",
            height=18,
            padx=10,
            pady=10,
        )
        self.log_box.pack(fill="both", expand=True)
        self.log_box.configure(state="disabled")

        status = ttk.Label(self.root, textvariable=self.status_var, style="Status.TLabel", padding=(16, 0, 16, 12))
        status.pack(fill="x")

    def append_log(self, text):
        def _write():
            self.log_box.configure(state="normal")
            self.log_box.insert("end", f"{text}\n")
            self.log_box.see("end")
            self.log_box.configure(state="disabled")
            self.status_var.set(text)

        self.root.after(0, _write)

    def pick_pdf(self):
        path = filedialog.askopenfilename(
            title="Selecciona un PDF",
            filetypes=[("PDF", "*.pdf"), ("Tots", "*.*")],
        )
        if path:
            self.source_pdf = Path(path)
            self.pdf_path_var.set(path)
            self.append_log("PDF seleccionat.")

    def open_ocr_pdf(self):
        if not self.ocr_pdf or not self.ocr_pdf.exists():
            messagebox.showinfo(APP_NAME, "Encara no hi ha cap PDF OCR generat.")
            return
        os.startfile(self.ocr_pdf)  # type: ignore[attr-defined]

    def load_pages_into_ui(self):
        for item in self.page_tree.get_children():
            self.page_tree.delete(item)
        for page in self.pages:
            preview = page.text.splitlines()[0] if page.text.strip() else "(sense text)"
            self.page_tree.insert("", "end", iid=str(page.number - 1), values=(page.number, preview[:48]))
        if self.pages:
            self.page_tree.selection_set("0")
            self.page_tree.focus("0")
            self.populate_editor(0)

    def populate_editor(self, index: int):
        self.current_page_index = index
        self.suspend_dirty = True
        try:
            self.editor.delete("1.0", "end")
            self.editor.insert("1.0", self.pages[index].text)
            self.editor.edit_modified(False)
        finally:
            self.suspend_dirty = False
        self.dirty = False

    def on_page_select(self, _event):
        selection = self.page_tree.selection()
        if not selection:
            return
        target = int(selection[0])
        if self.current_page_index is not None and target != self.current_page_index and self.dirty:
            self.apply_current_page()
        self.populate_editor(target)

    def on_editor_modified(self, _event):
        if self.suspend_dirty:
            self.editor.edit_modified(False)
            return
        self.dirty = True
        self.status_var.set("Hi ha canvis pendents d'aplicar en la pàgina actual.")
        self.editor.edit_modified(False)

    def apply_current_page(self):
        if self.current_page_index is None or not self.pages:
            return True
        self.pages[self.current_page_index].text = self.editor.get("1.0", "end-1c")
        preview = self.pages[self.current_page_index].text.splitlines()[0] if self.pages[self.current_page_index].text.strip() else "(sense text)"
        self.page_tree.item(str(self.current_page_index), values=(self.pages[self.current_page_index].number, preview[:48]))
        self.dirty = False
        self.status_var.set("Canvis aplicats.")
        return True

    def process_ocr(self):
        if self.worker_thread and self.worker_thread.is_alive():
            messagebox.showinfo(APP_NAME, "Ja hi ha un procés OCR en marxa.")
            return
        if not self.source_pdf or not self.source_pdf.exists():
            messagebox.showwarning(APP_NAME, "Selecciona abans un PDF.")
            return

        def worker():
            try:
                client = PDF24OCRClient(self.append_log)
                ocr_pdf = client.ocr_pdf(self.source_pdf)
                self.ocr_pdf = ocr_pdf
                self.root.after(0, lambda: self.ocr_path_var.set(str(ocr_pdf)))
                self.append_log("Extraient text del PDF OCR...")
                pages = extract_document_pages(ocr_pdf)
                self.pages = pages
                self.root.after(0, self.load_pages_into_ui)
                self.append_log(f"OCR completat. {len(pages)} pàgines carregades a l'editor.")
            except Exception as exc:
                self.append_log(f"Error: {exc}")
                self.root.after(0, lambda: messagebox.showerror(APP_NAME, str(exc)))

        self.worker_thread = threading.Thread(target=worker, daemon=True)
        self.worker_thread.start()
        self.append_log("Iniciant procés OCR...")

    def reload_text_from_ocr(self):
        if not self.ocr_pdf or not self.ocr_pdf.exists():
            messagebox.showinfo(APP_NAME, "Encara no hi ha cap PDF OCR per a reprocessar.")
            return
        self.apply_current_page()
        try:
            self.pages = extract_document_pages(self.ocr_pdf)
            self.load_pages_into_ui()
            self.append_log("Text reextret des del PDF OCR.")
        except Exception as exc:
            messagebox.showerror(APP_NAME, str(exc))

    def export_to_docx(self):
        if not self.pages:
            messagebox.showwarning(APP_NAME, "No hi ha text carregat per a exportar.")
            return
        self.apply_current_page()
        output = filedialog.asksaveasfilename(
            title="Guardar DOCX",
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
        )
        if not output:
            return
        export_docx(self.pages, Path(output))
        self.append_log(f"DOCX exportat: {output}")

    def export_to_pdf(self):
        if not self.pages:
            messagebox.showwarning(APP_NAME, "No hi ha text carregat per a exportar.")
            return
        self.apply_current_page()
        output = filedialog.asksaveasfilename(
            title="Guardar PDF net",
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
        )
        if not output:
            return
        export_clean_pdf(self.pages, Path(output))
        self.append_log(f"PDF net exportat: {output}")


def main():
    root = tk.Tk()
    if ICON_PATH.exists():
        try:
            root.iconbitmap(str(ICON_PATH))
        except Exception:
            pass
    AutoOCR24App(root)
    root.mainloop()


if __name__ == "__main__":
    main()
