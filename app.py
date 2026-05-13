import datetime as dt
import html
import json
import os
import re
import shutil
import statistics
import subprocess
import sys
import tempfile
import threading
import tkinter as tk
import time
import unicodedata
import urllib.parse
import urllib.request
import urllib.error
import webbrowser
from dataclasses import asdict, dataclass
from pathlib import Path
from tkinter import filedialog, messagebox, ttk

import openpyxl
import pdfplumber
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from PIL import Image, ImageTk
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer


APP_NAME = "AutoCPV"
APP_VERSION = "1.5"
DEFAULT_FORM_URL = "https://docs.google.com/forms/d/e/1FAIpQLScdmuAuYu918Iv28w3v94kjs_uW2vyRSOAubcrnaWIyQTuQXA/viewform"
PDF24_OCR_EXE = Path(r"C:\Program Files\PDF24\pdf24-Ocr.exe")
PDF24_WORKING_DIR = Path(r"C:\Program Files\PDF24")
PROJECT_DIR = Path(r"C:\Users\solso\Documents\New project")
PROMPT_FILENAME = "PROMPT AutoCPV.txt"
LOGO_FILENAME = "logo-trimmed.png"
NVIDIA_CHAT_COMPLETIONS_URL = "https://integrate.api.nvidia.com/v1/chat/completions"
DEFAULT_NVIDIA_MODEL = "openai/gpt-oss-120b"
DEFAULT_NVIDIA_MAX_TOKENS = 16000
CONFIG_DIR = Path(os.environ.get("APPDATA", PROJECT_DIR)) / APP_NAME
CONFIG_PATH = CONFIG_DIR / "settings.json"
SESSION_FILETYPES = [("AutoCPV Session", "*.autocpv.json"), ("JSON", "*.json")]
DEFAULT_FACEBOOK_SEARCH = "https://www.facebook.com/search/top?q="
LOGO_PATH = PROJECT_DIR / "assets" / LOGO_FILENAME


def app_search_dirs() -> list[Path]:
    dirs = []
    if getattr(sys, "frozen", False):
        exe_dir = Path(sys.executable).resolve().parent
        dirs.extend([exe_dir, exe_dir.parent])
    bundle_root = getattr(sys, "_MEIPASS", "")
    if bundle_root:
        bundle_dir = Path(bundle_root)
        dirs.extend([bundle_dir, bundle_dir / "assets"])
    dirs.extend([PROJECT_DIR, Path(__file__).resolve().parent, Path.cwd()])
    unique_dirs = []
    for folder in dirs:
        if folder not in unique_dirs:
            unique_dirs.append(folder)
    return unique_dirs


def resolve_prompt_path() -> Path:
    for folder in app_search_dirs():
        candidate = folder / PROMPT_FILENAME
        if candidate.exists():
            return candidate
    return PROJECT_DIR / PROMPT_FILENAME


def resolve_asset_path(filename: str) -> Path:
    for folder in app_search_dirs():
        candidates = [folder / filename, folder / "assets" / filename]
        for candidate in candidates:
            if candidate.exists():
                return candidate
    return PROJECT_DIR / "assets" / filename


def resolve_pdf24_ocr_exe() -> Path:
    candidates = [
        PDF24_OCR_EXE,
        Path(os.environ.get("ProgramFiles", r"C:\Program Files")) / "PDF24" / "pdf24-Ocr.exe",
        Path(os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)")) / "PDF24" / "pdf24-Ocr.exe",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return PDF24_OCR_EXE


def load_app_config() -> dict:
    if not CONFIG_PATH.exists():
        return {}
    try:
        data = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except (OSError, json.JSONDecodeError):
        return {}


def save_app_config(config: dict):
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    CONFIG_PATH.write_text(json.dumps(config, indent=2, ensure_ascii=False), encoding="utf-8")


FIELD_LABELS = {
    "localitat": "Localitat",
    "data": "Data",
    "categoria": "Categoria",
    "altres": "En cas d'altres, quina?",
    "nom": "Nom de l'activitat",
    "companyia": "Companyia, artista",
    "lloc": "Lloc",
    "llengua": "Llengua de l'activitat",
    "preu": "Preu",
    "regidoria": "Regidoria organitzadora",
    "publicitat": "Llengua de la publicitat",
    "font": "Font",
    "persona": "Persona que ha introduit les dades",
}

EXPECTED_HEADERS = {
    "Localitat": "localitat",
    "Data": "data",
    "Categoria": "categoria",
    "En cas d'altres, quina?": "altres",
    "Nom de l'activitat": "nom",
    "Companyia, artista": "companyia",
    "Lloc": "lloc",
    "Llengua de l'activitat": "llengua",
    "Preu": "preu",
    "Regidoria organitzadora": "regidoria",
    "Llengua de la publicitat": "publicitat",
    "Llengua de la publicitat ⚠️": "publicitat",
    "Font": "font",
    "Persona que ha introduit les dades": "persona",
    "Persona que ha introduït les dades": "persona",
}

FALLBACK_FIELD_IDS = {
    "localitat": "2074597570",
    "data": "1851780647",
    "categoria": "25855538",
    "altres": "30622121",
    "nom": "757267152",
    "companyia": "799207947",
    "lloc": "319720355",
    "llengua": "2015577411",
    "preu": "1027978183",
    "regidoria": "936300587",
    "publicitat": "744125401",
    "font": "54121299",
    "persona": "1706428244",
}

PERSON_OPTIONS = ["Arantxa", "Erika", "Maria", "Maria R.", "Pol"]
CATEGORIA_OPTIONS = [
    "Teatre",
    "Cinema",
    "Música",
    "Exposició",
    "Presentació llibre",
    "Lectura poemes",
    "Monòlegs",
    "Contacontes",
    "Taller de teatre",
    "Taller literari",
    "Dansa",
    "Taller/curs de dansa",
    "Activitats sobre patrimoni",
    "Lectura en veu alta",
    "Club de lectura",
    "Conferència",
    "Altres",
]
LLENGUA_ACTIVITY_OPTIONS = [
    "Valencià/català",
    "Espanyol",
    "Anglès",
    "Francès",
    "Activitat sense llengua",
    "No hi ha informació",
]
REGIDORIA_OPTIONS = [
    "Cultura",
    "Joventut",
    "Igualtat",
    "Altres",
    "En col·laboració o organitzat per una entitat ciutadana",
]
PUBLICITAT_OPTIONS = [
    "Valencià/català",
    "Espanyol",
    "Bilingüe",
    "Anglès",
    "Altres",
]

FILTER_OPTIONS = {
    "Totes": "all",
    "Pendents": "pending",
    "Enviades": "sent",
    "Amb error": "error",
    "Per revisar": "invalid",
}

REQUIRED_FIELDS = [
    "localitat",
    "data",
    "categoria",
    "nom",
    "llengua",
    "regidoria",
    "publicitat",
    "font",
    "persona",
]

FIELD_OPTIONS = {
    "categoria": CATEGORIA_OPTIONS,
    "llengua": LLENGUA_ACTIVITY_OPTIONS,
    "regidoria": REGIDORIA_OPTIONS,
    "publicitat": PUBLICITAT_OPTIONS,
    "persona": PERSON_OPTIONS,
}

MASS_EDIT_FIELDS = {
    "Llengua de l'activitat": "llengua",
    "Categoria": "categoria",
    "Regidoria organitzadora": "regidoria",
    "Llengua de la publicitat": "publicitat",
    "Persona": "persona",
}

LONG_TEXT_FIELDS = {"nom", "companyia", "lloc", "font", "altres"}

COLORS = {
    "red": "#C8102E",
    "red_dark": "#9E0E24",
    "cream": "#F6F0E8",
    "paper": "#FFFDF9",
    "charcoal": "#1F1F1F",
    "muted": "#655D57",
    "line": "#E5D8CB",
    "rose": "#F7E4E8",
    "ok": "#1C7C54",
    "warning": "#C4831B",
    "warning_bg": "#FFF3D6",
    "error_bg": "#F9E2E4",
    "sent_bg": "#E5F4EC",
}


@dataclass
class Record:
    localitat: str = ""
    data: str = ""
    categoria: str = ""
    altres: str = ""
    nom: str = ""
    companyia: str = ""
    lloc: str = ""
    llengua: str = ""
    preu: str = ""
    regidoria: str = ""
    publicitat: str = ""
    font: str = ""
    persona: str = "Pol"
    status: str = "Pendent"
    status_detail: str = ""


@dataclass
class OCRPage:
    number: int
    text: str


class OCRCancelled(Exception):
    pass


def structured_page_text(page) -> str:
    words = page.extract_words(x_tolerance=2, y_tolerance=3, keep_blank_chars=False)
    if not words:
        return (page.extract_text() or "").strip()

    words = sorted(words, key=lambda item: (round(item["top"], 1), item["x0"]))
    lines = []
    for word in words:
        top = word["top"]
        bottom = word["bottom"]
        height = bottom - top
        if not lines:
            lines.append({"top": top, "bottom": bottom, "words": [word]})
            continue
        previous = lines[-1]
        tolerance = max(3.0, min(height, previous["bottom"] - previous["top"]) * 0.55)
        if abs(top - previous["top"]) <= tolerance:
            previous["words"].append(word)
            previous["bottom"] = max(previous["bottom"], bottom)
        else:
            lines.append({"top": top, "bottom": bottom, "words": [word]})

    text_lines = []
    heights = [line["bottom"] - line["top"] for line in lines]
    median_height = statistics.median(heights) if heights else 12
    previous_bottom = None

    for line in lines:
        ordered_words = sorted(line["words"], key=lambda item: item["x0"])
        average_char_widths = [
            (item["x1"] - item["x0"]) / max(len(item["text"]), 1)
            for item in ordered_words
            if item["text"]
        ]
        average_char_width = statistics.mean(average_char_widths) if average_char_widths else 6

        chunks = []
        previous_word = None
        for word in ordered_words:
            if previous_word is not None:
                gap = word["x0"] - previous_word["x1"]
                chunks.append("    " if gap > average_char_width * 5 else " ")
            chunks.append(word["text"])
            previous_word = word

        if previous_bottom is not None and line["top"] - previous_bottom > median_height * 0.9:
            text_lines.append("")
        text_lines.append("".join(chunks).strip())
        previous_bottom = line["bottom"]

    return "\n".join(text_lines).strip()


def extract_document_pages(pdf_path: Path) -> list[OCRPage]:
    pages = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            pages.append(OCRPage(number=page_number, text=structured_page_text(page)))
    return pages


def export_docx(pages: list[OCRPage], output_path: Path):
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    for page_index, page in enumerate(pages):
        blocks = [block.strip() for block in page.text.split("\n\n")]
        for block in blocks:
            paragraph = document.add_paragraph()
            lines = [line.rstrip() for line in block.splitlines() if line.strip() or len(blocks) == 1]
            for index, line in enumerate(lines):
                run = paragraph.add_run(line)
                if index < len(lines) - 1:
                    run.add_break(WD_BREAK.LINE)
        if page_index < len(pages) - 1:
            document.add_page_break()

    document.save(output_path)


def export_clean_pdf(pages: list[OCRPage], output_path: Path):
    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "CleanText",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=11,
        leading=15,
        spaceAfter=8,
    )
    story = []
    for page_index, page in enumerate(pages):
        blocks = [block.strip() for block in page.text.split("\n\n") if block.strip()]
        for block in blocks:
            escaped = "<br/>".join(html.escape(line) for line in block.splitlines())
            story.append(Paragraph(escaped, normal))
            story.append(Spacer(1, 6))
        if page_index < len(pages) - 1:
            story.append(PageBreak())

    document = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=52,
        rightMargin=52,
        topMargin=52,
        bottomMargin=52,
    )
    document.build(story)


def normalize_label(text: str) -> str:
    text = text.replace("\xa0", " ").strip()
    return re.sub(r"\s+", " ", text)


def normalize_form_key(text: str) -> str:
    text = html.unescape(text or "")
    text = normalize_label(text).casefold()
    text = "".join(ch for ch in unicodedata.normalize("NFKD", text) if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", "", text)


def find_header_row(worksheet):
    for row_idx in range(1, min(worksheet.max_row, 30) + 1):
        headers = [normalize_label(str(cell.value or "")) for cell in worksheet[row_idx]]
        matches = sum(1 for item in headers if item in EXPECTED_HEADERS)
        if matches >= 5:
            return row_idx, headers
    raise ValueError("No he trobat una fila de capçaleres recognoscible en l'Excel.")


def excel_date_to_iso(value) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, dt.datetime):
        return value.date().isoformat()
    if isinstance(value, dt.date):
        return value.isoformat()
    text = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return dt.datetime.strptime(text, fmt).date().isoformat()
        except ValueError:
            pass
    return text


def normalize_price(value) -> str:
    if value in (None, ""):
        return ""
    if isinstance(value, (int, float)):
        return str(int(value)) if float(value).is_integer() else str(value)
    text = str(value).strip().replace("€", "").replace(",", ".")
    return re.sub(r"\s+", "", text)


def is_valid_date(value: str) -> bool:
    return bool(re.fullmatch(r"\d{4}-\d{2}-\d{2}", value or ""))


def is_valid_price(value: str) -> bool:
    return bool(re.fullmatch(r"\d+(?:\.\d+)?", value or ""))


def load_excel_records(path: str, default_person: str, fallback_font: str):
    workbook = openpyxl.load_workbook(path, data_only=True)
    worksheet = workbook[workbook.sheetnames[0]]
    header_row, headers = find_header_row(worksheet)

    mapping = {}
    for idx, header in enumerate(headers):
        key = EXPECTED_HEADERS.get(header)
        if key:
            mapping[key] = idx

    records = []
    for row in worksheet.iter_rows(min_row=header_row + 1, values_only=True):
        if not any(value not in (None, "") for value in row):
            continue
        record = Record(persona=default_person.strip() or "Pol")
        for key, col_idx in mapping.items():
            if col_idx >= len(row):
                continue
            value = row[col_idx]
            if key == "data":
                setattr(record, key, excel_date_to_iso(value))
            elif key == "preu":
                setattr(record, key, normalize_price(value))
            else:
                setattr(record, key, "" if value is None else str(value).strip())
        if not record.font:
            record.font = fallback_font.strip()
        if not record.persona:
            record.persona = default_person.strip() or "Pol"
        records.append(record)
    return records


def autocpv_excel_headers():
    return [FIELD_LABELS[key] for key in FIELD_LABELS]


def pages_to_prompt_text(pages: list[OCRPage]) -> str:
    chunks = []
    for page in pages:
        chunks.append(f"--- Pàgina {page.number} ---\n{page.text.strip()}")
    return "\n\n".join(chunks).strip()


def build_autocpv_json_schema():
    row_properties = {key: {"type": "string", "description": FIELD_LABELS[key]} for key in FIELD_LABELS}
    return {
        "type": "object",
        "additionalProperties": False,
        "required": ["rows", "report"],
        "properties": {
            "rows": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "required": list(FIELD_LABELS.keys()),
                    "properties": row_properties,
                },
            },
            "report": {
                "type": "object",
                "additionalProperties": False,
                "required": ["localitat", "pages_reviewed", "total_pages", "activities_detected", "rows_generated", "discarded"],
                "properties": {
                    "localitat": {"type": "string"},
                    "pages_reviewed": {"type": "integer"},
                    "total_pages": {"type": "integer"},
                    "activities_detected": {"type": "integer"},
                    "rows_generated": {"type": "integer"},
                    "discarded": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "additionalProperties": False,
                            "required": ["name", "reason", "page"],
                            "properties": {
                                "name": {"type": "string"},
                                "reason": {"type": "string"},
                                "page": {"type": "string"},
                            },
                        },
                    },
                },
            },
        },
    }


def response_output_text(payload: dict) -> str:
    if payload.get("output_text"):
        return payload["output_text"]
    parts = []
    for item in payload.get("output", []):
        for content in item.get("content", []):
            text = content.get("text")
            if text:
                parts.append(text)
    return "\n".join(parts).strip()


def chat_completion_text(payload: dict) -> str:
    choices = payload.get("choices", [])
    if not choices:
        return ""
    message = choices[0].get("message", {})
    content = message.get("content", "")
    if isinstance(content, list):
        return "".join(part.get("text", "") for part in content if isinstance(part, dict))
    return str(content or "").strip()


def parse_json_response(text: str) -> dict:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s*```$", "", cleaned)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start != -1 and end != -1 and end > start:
            return json.loads(cleaned[start : end + 1])
        raise


def post_nvidia_chat(messages: list[dict], api_key: str, model: str, max_tokens: int, temperature: float = 0.1) -> dict:
    request_payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "top_p": 1,
        "max_tokens": max_tokens,
        "response_format": {"type": "json_object"},
    }
    request = urllib.request.Request(
        NVIDIA_CHAT_COMPLETIONS_URL,
        data=json.dumps(request_payload).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=420) as response:
            return json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"NVIDIA ha retornat un error HTTP {exc.code}: {detail}") from exc
    except urllib.error.URLError as exc:
        raise RuntimeError(f"No s'ha pogut connectar amb NVIDIA: {exc}") from exc


def repair_json_with_nvidia(raw_text: str, api_key: str, model: str, max_tokens: int, progress_callback=None) -> dict:
    if progress_callback:
        progress_callback("La resposta no era JSON valid. Intentant reparar-la automaticament...")
    schema_text = json.dumps(build_autocpv_json_schema(), ensure_ascii=False)
    messages = [
        {
            "role": "system",
            "content": (
                "Eres un reparador de JSON. Devuelve exclusivamente un objeto JSON valido. "
                "No anadas Markdown, comentarios ni explicaciones."
            ),
        },
        {
            "role": "user",
            "content": (
                "Repara el siguiente texto para que sea JSON valido y cumpla este esquema. "
                "Conserva todas las filas de actividades que puedas sin inventar datos nuevos.\n\n"
                f"ESQUEMA:\n{schema_text}\n\n"
                f"TEXTO A REPARAR:\n{raw_text}"
            ),
        },
    ]
    repaired_payload = post_nvidia_chat(messages, api_key, model, max_tokens, temperature=0)
    repaired_text = chat_completion_text(repaired_payload)
    if not repaired_text:
        raise RuntimeError("NVIDIA no ha retornat text en l'intent de reparacio JSON.")
    return parse_json_response(repaired_text)


def call_ai_for_autocpv(prompt_text: str, ocr_text: str, api_key: str, progress_callback=None, model: str | None = None, max_tokens: int | None = None):
    model = model or os.environ.get("NVIDIA_MODEL", DEFAULT_NVIDIA_MODEL)
    max_tokens = max_tokens or int(os.environ.get("NVIDIA_MAX_TOKENS", str(DEFAULT_NVIDIA_MAX_TOKENS)))
    system_text = (
        "Eres un extractor de actividades culturales para AutoCPV. "
        "Devuelve exclusivamente JSON válido que cumpla el esquema. "
        "No escribas Markdown ni explicaciones fuera del JSON. "
        "Antes de finalizar, comprueba que el JSON se puede parsear con json.loads."
    )
    integration_text = (
        "INSTRUCCIÓN DE INTEGRACIÓN: aunque el prompt pida crear un Excel, en esta aplicación debes devolver JSON "
        "con las claves exactas del esquema. AutoCPV creará el Excel después. "
        "Las fechas deben ir en formato YYYY-MM-DD para que AutoCPV las convierta en fechas reales de Excel."
    )
    schema_text = json.dumps(build_autocpv_json_schema(), ensure_ascii=False)
    user_text = (
        f"{prompt_text}\n\n{integration_text}\n\n"
        f"ESQUEMA JSON OBLIGATORIO:\n{schema_text}\n\n"
        f"TEXT OCR DEL PDF:\n{ocr_text}"
    )
    request_payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": system_text},
            {"role": "user", "content": user_text},
        ],
        "temperature": 0.2,
        "top_p": 1,
        "max_tokens": max_tokens,
        "response_format": {"type": "json_object"},
    }
    if progress_callback:
        progress_callback(f"Enviant text OCR a NVIDIA ({model}, max_tokens={max_tokens})...")
    response_payload = post_nvidia_chat(request_payload["messages"], api_key, model, max_tokens, temperature=0.2)

    output_text = chat_completion_text(response_payload)
    if not output_text:
        raise RuntimeError("NVIDIA no ha retornat text estructurat.")
    try:
        return parse_json_response(output_text)
    except json.JSONDecodeError as exc:
        try:
            return repair_json_with_nvidia(output_text, api_key, model, max_tokens, progress_callback)
        except json.JSONDecodeError as repair_exc:
            raise RuntimeError(f"La resposta de NVIDIA no es JSON valid: {repair_exc}") from exc


def coerce_excel_date(value):
    if not value:
        return ""
    if isinstance(value, dt.datetime):
        return value.date()
    if isinstance(value, dt.date):
        return value
    text = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y"):
        try:
            return dt.datetime.strptime(text, fmt).date()
        except ValueError:
            pass
    return text


def write_autocpv_excel(rows: list[dict], output_path: Path):
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = "Activitats"
    worksheet.append(autocpv_excel_headers())
    for row_data in rows:
        values = []
        for key in FIELD_LABELS:
            value = row_data.get(key, "")
            if key == "data":
                value = coerce_excel_date(value)
            values.append(value)
        worksheet.append(values)
    for cell in worksheet[1]:
        cell.font = openpyxl.styles.Font(bold=True)
    date_col = list(FIELD_LABELS.keys()).index("data") + 1
    for row in worksheet.iter_rows(min_row=2, min_col=date_col, max_col=date_col):
        for cell in row:
            if isinstance(cell.value, dt.date):
                cell.number_format = "dd/mm/yyyy"
    for column_cells in worksheet.columns:
        max_length = max(len(str(cell.value or "")) for cell in column_cells)
        worksheet.column_dimensions[column_cells[0].column_letter].width = min(max(max_length + 2, 12), 42)
    workbook.save(output_path)


def extract_form_metadata(form_url: str):
    view_url = form_url.strip()
    if "formResponse" in view_url:
        view_url = view_url.replace("formResponse", "viewform")
    response_url = view_url.replace("viewform", "formResponse")

    request = urllib.request.Request(view_url, headers={"User-Agent": "Mozilla/5.0"})
    with urllib.request.urlopen(request, timeout=30) as response:
        html_text = response.read().decode("utf-8", errors="ignore")

    field_ids = FALLBACK_FIELD_IDS.copy()
    normalized_field_labels = {normalize_form_key(label): key for key, label in FIELD_LABELS.items()}
    data_params_pattern = re.compile(r'data-params="%\.\@\.\[(\d+),&quot;(.*?)&quot;,null,\d+,\[\[(\d+),', re.DOTALL)
    for match in data_params_pattern.finditer(html_text):
        raw_label = match.group(2)
        field_id = match.group(3)
        normalized_label = normalize_form_key(raw_label)
        key = normalized_field_labels.get(normalized_label)
        if key:
            field_ids[key] = field_id

    form_action_match = re.search(r'<form[^>]*action="([^"]+)"', html_text)
    if form_action_match:
        response_url = html.unescape(form_action_match.group(1))

    fbzx_match = re.search(r'name="fbzx"\s+value="([^"]+)"', html_text)
    partial_match = re.search(r'name="partialResponse"\s+value="([^"]+)"', html_text)
    page_history_match = re.search(r'name="pageHistory"\s+value="([^"]+)"', html_text)
    fvv_match = re.search(r'name="fvv"\s+value="([^"]+)"', html_text)

    return {
        "response_url": response_url,
        "fbzx": html.unescape(fbzx_match.group(1)) if fbzx_match else "",
        "partialResponse": html.unescape(partial_match.group(1)) if partial_match else "",
        "pageHistory": html.unescape(page_history_match.group(1)) if page_history_match else "0",
        "fvv": html.unescape(fvv_match.group(1)) if fvv_match else "1",
        "field_ids": field_ids,
    }


def build_payload(record: Record, metadata):
    field_ids = metadata["field_ids"]
    payload = {
        f'entry.{field_ids["localitat"]}': record.localitat,
        f'entry.{field_ids["data"]}_year': record.data[0:4],
        f'entry.{field_ids["data"]}_month': str(int(record.data[5:7])),
        f'entry.{field_ids["data"]}_day': str(int(record.data[8:10])),
        f'entry.{field_ids["categoria"]}': record.categoria,
        f'entry.{field_ids["nom"]}': record.nom,
        f'entry.{field_ids["companyia"]}': record.companyia,
        f'entry.{field_ids["lloc"]}': record.lloc,
        f'entry.{field_ids["llengua"]}': record.llengua,
        f'entry.{field_ids["regidoria"]}': record.regidoria,
        f'entry.{field_ids["publicitat"]}': record.publicitat,
        f'entry.{field_ids["font"]}': record.font,
        f'entry.{field_ids["persona"]}': record.persona,
        "fvv": metadata["fvv"] or "1",
        "pageHistory": metadata["pageHistory"] or "0",
        "fbzx": metadata["fbzx"],
        "submissionTimestamp": "-1",
    }
    if metadata["partialResponse"]:
        payload["partialResponse"] = metadata["partialResponse"]
    if record.altres:
        payload[f'entry.{field_ids["altres"]}'] = record.altres
    if record.preu != "":
        payload[f'entry.{field_ids["preu"]}'] = record.preu
    return payload


class FormFillerApp:
    def __init__(self, root):
        self.root = root
        self.root.title(APP_NAME)
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        window_width = min(1540, max(1100, screen_width - 80))
        window_height = min(980, max(720, screen_height - 100))
        self.root.geometry(f"{window_width}x{window_height}")
        self.root.minsize(1050, 680)
        self.root.configure(bg=COLORS["cream"])

        self.records = []
        self.form_metadata = None
        self.current_index = None
        self.visible_indices = []
        self.worker = None
        self.dirty = False
        self.suspend_dirty = False
        self.autosave_after_id = None
        self.logo_image = None
        self.last_deleted = None
        self.history_entries = []
        self.editor_widgets = {}
        self.help_popup = None
        self.closing = False
        self.after_ids = set()
        self.source_pdf = None
        self.ocr_pdf = None
        self.ocr_pages = []
        self.ocr_current_page_index = None
        self.ocr_dirty = False
        self.ocr_suspend_dirty = False
        self.pdf24_process = None
        self.config_data = load_app_config()
        ai_config = self.config_data.get("ai", {}) if isinstance(self.config_data.get("ai", {}), dict) else {}

        self.excel_path_var = tk.StringVar()
        self.form_url_var = tk.StringVar(value=DEFAULT_FORM_URL)
        self.person_var = tk.StringVar(value="Pol")
        self.fallback_font_var = tk.StringVar(value="")
        self.status_filter_var = tk.StringVar(value="Totes")
        self.mass_edit_field_var = tk.StringVar(value="Llengua de l'activitat")
        self.mass_edit_value_var = tk.StringVar(value=LLENGUA_ACTIVITY_OPTIONS[0])
        self.status_var = tk.StringVar(value="A punt.")
        self.validation_var = tk.StringVar(value="Sense validacions pendents.")
        self.editor_vars = {key: tk.StringVar() for key in FIELD_LABELS}
        self.ocr_pdf_path_var = tk.StringVar()
        self.ocr_result_path_var = tk.StringVar(value="Encara no s'ha generat.")
        self.ocr_status_var = tk.StringVar(value="A punt per a processar PDFs.")
        self.ocr_summary_var = tk.StringVar(value="Selecciona un PDF per començar.")
        self.ocr_progress_var = tk.DoubleVar(value=0.0)
        self.ocr_localitat_var = tk.StringVar()
        self.ocr_font_var = tk.StringVar()
        self.nvidia_api_key_var = tk.StringVar(value=ai_config.get("nvidia_api_key") or os.environ.get("NVIDIA_API_KEY", ""))
        self.nvidia_model_var = tk.StringVar(value=ai_config.get("nvidia_model") or os.environ.get("NVIDIA_MODEL", DEFAULT_NVIDIA_MODEL))
        self.nvidia_max_tokens_var = tk.StringVar(value=str(ai_config.get("nvidia_max_tokens") or os.environ.get("NVIDIA_MAX_TOKENS", DEFAULT_NVIDIA_MAX_TOKENS)))

        self.configure_styles()
        self.build_ui()
        self.bind_editor_events()
        self.bind_shortcuts()
        self.root.protocol("WM_DELETE_WINDOW", self.on_close)
        self.show_splash()

    def configure_styles(self):
        style = ttk.Style(self.root)
        if "vista" in style.theme_names():
            style.theme_use("vista")

        style.configure("Root.TFrame", background=COLORS["cream"])
        style.configure("Card.TFrame", background=COLORS["paper"], relief="flat")
        style.configure("Header.TFrame", background=COLORS["red"])
        style.configure("HeaderTitle.TLabel", background=COLORS["red"], foreground="white", font=("Segoe UI Semibold", 22))
        style.configure("HeaderSub.TLabel", background=COLORS["red"], foreground="white", font=("Segoe UI", 10))
        style.configure("Section.TLabel", background=COLORS["paper"], foreground=COLORS["charcoal"], font=("Segoe UI Semibold", 10))
        style.configure("Body.TLabel", background=COLORS["paper"], foreground=COLORS["charcoal"], font=("Segoe UI", 10))
        style.configure("Status.TLabel", background=COLORS["cream"], foreground=COLORS["charcoal"], font=("Segoe UI Semibold", 10))
        style.configure("Neutral.TButton", font=("Segoe UI Semibold", 10))
        style.configure("Accent.Horizontal.TProgressbar", troughcolor=COLORS["cream"], background=COLORS["red"], bordercolor=COLORS["line"], lightcolor=COLORS["red"], darkcolor=COLORS["red"])
        style.configure("Treeview", rowheight=28, font=("Segoe UI", 10), fieldbackground="white", background="white", foreground=COLORS["charcoal"])
        style.configure("Treeview.Heading", font=("Segoe UI Semibold", 10))
        style.map("Treeview", background=[("selected", COLORS["red"])], foreground=[("selected", "white")])

    def build_ui(self):
        header = ttk.Frame(self.root, style="Header.TFrame", padding=(20, 18))
        header.pack(fill="x")
        header_left = ttk.Frame(header, style="Header.TFrame")
        header_left.pack(side="left", fill="x", expand=True)
        header_button_row = ttk.Frame(header_left, style="Header.TFrame")
        header_button_row.pack(anchor="w", pady=(0, 10))
        self.make_help_button(header_button_row, "Tecles ràpides", self.show_shortcuts_help, "Mostra la llista d'accessos ràpids.").pack(side="left")
        ttk.Label(header_left, text=APP_NAME, style="HeaderTitle.TLabel").pack(anchor="w")
        ttk.Label(header_left, text=f"Versió {APP_VERSION}", style="HeaderSub.TLabel").pack(anchor="w", pady=(2, 0))

        notebook_wrap = ttk.Frame(self.root, style="Root.TFrame", padding=(16, 14, 16, 8))
        notebook_wrap.pack(fill="both", expand=True)
        self.notebook = ttk.Notebook(notebook_wrap)
        self.notebook.pack(fill="both", expand=True)
        self.main_tab = ttk.Frame(self.notebook, style="Root.TFrame")
        self.ocr_tab = ttk.Frame(self.notebook, style="Root.TFrame")
        self.notebook.add(self.main_tab, text="Formularis")
        self.notebook.add(self.ocr_tab, text="OCR de PDFs")

        top = ttk.Frame(self.main_tab, style="Root.TFrame", padding=(0, 0, 0, 8))
        top.pack(fill="x")
        top_card = ttk.Frame(top, style="Card.TFrame", padding=14)
        top_card.pack(fill="x")

        ttk.Label(top_card, text="Excel", style="Section.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Entry(top_card, textvariable=self.excel_path_var, width=92).grid(row=0, column=1, sticky="ew", padx=8, pady=4)
        self.make_help_button(top_card, "Buscar", self.pick_excel, "Selecciona un fitxer Excel del teu ordinador.").grid(row=0, column=2, padx=4)
        self.make_help_button(top_card, "Carregar Excel", self.load_excel, "Llig l'Excel i carrega les files a la taula.").grid(row=0, column=3, padx=4)
        self.make_help_button(top_card, "Guardar sessió", self.save_session, "Guarda l'estat actual per continuar després.").grid(row=0, column=4, padx=4)
        self.make_help_button(top_card, "Obrir sessió", self.load_session, "Recupera una sessió guardada anteriorment.").grid(row=0, column=5, padx=4)

        ttk.Label(top_card, text="Formulari", style="Section.TLabel").grid(row=1, column=0, sticky="w")
        ttk.Entry(top_card, textvariable=self.form_url_var, width=92).grid(row=1, column=1, sticky="ew", padx=8, pady=4)
        self.make_help_button(top_card, "Llegir formulari", self.load_form, "Detecta els camps del Google Form automàticament.").grid(row=1, column=2, padx=4)
        self.make_help_button(top_card, "Previsualitzar enviament", self.preview_current_payload, "Mostra exactament què s'enviarà al formulari.").grid(row=1, column=3, padx=4)
        self.make_help_button(top_card, "Revisió ampla", self.open_review_mode, "Obri un editor més gran per revisar la fila.").grid(row=1, column=4, padx=4)

        ttk.Label(top_card, text="Persona", style="Section.TLabel").grid(row=2, column=0, sticky="w")
        ttk.Combobox(top_card, textvariable=self.person_var, values=PERSON_OPTIONS, width=16, state="readonly").grid(row=2, column=1, sticky="w", padx=8, pady=4)
        ttk.Label(top_card, text="Font per defecte", style="Section.TLabel").grid(row=2, column=1, sticky="e")
        ttk.Entry(top_card, textvariable=self.fallback_font_var, width=44).grid(row=2, column=2, sticky="ew", padx=8, pady=4)
        self.make_help_button(top_card, "Aplicar a totes", self.apply_fallback_font_to_all, "Copia la font per defecte a totes les files.").grid(row=2, column=3, padx=4)
        self.make_help_button(top_card, "Aplicar només a buides", self.apply_fallback_font_to_empty, "Ompli només les files que no tenen font.").grid(row=2, column=4, padx=4)
        top_card.columnconfigure(1, weight=1)

        controls = ttk.Frame(self.main_tab, style="Root.TFrame", padding=(0, 0, 0, 10))
        controls.pack(fill="x")
        controls_card = ttk.Frame(controls, style="Card.TFrame", padding=10)
        controls_card.pack(fill="x")
        self.make_help_button(controls_card, "Aplicar canvis", self.apply_current_record, "Guarda manualment els canvis de la fila actual.").pack(side="left", padx=4)
        self.make_help_button(controls_card, "Eliminar fila", self.delete_current_record, "Esborra completament les files seleccionades.").pack(side="left", padx=4)
        self.make_help_button(controls_card, "Desfer eliminació", self.undo_delete_record, "Recupera l'última eliminació de files.").pack(side="left", padx=4)
        self.make_help_button(controls_card, "Buscar a Google", self.open_google_search, "Busca l'activitat actual a Google.").pack(side="left", padx=4)
        self.make_help_button(controls_card, "Buscar font", self.open_source_helper, "Ajuda a trobar una font si encara no en tens.").pack(side="left", padx=4)
        self.make_help_button(controls_card, "Primera font", self.apply_google_first_result_to_selected, "Guarda la primera web de Google en la fila actual.").pack(side="left", padx=4)
        self.make_help_button(controls_card, "Primera font a totes", self.apply_google_first_result_to_all, "Prova d'omplir la font de totes les files amb el primer resultat.").pack(side="left", padx=4)
        self.make_help_button(controls_card, "Obrir font", self.open_source, "Obri l'enllaç de la font guardada.").pack(side="left", padx=4)
        self.make_help_button(controls_card, "Enviar fila seleccionada", self.submit_selected, "Envia només la fila actual al formulari.").pack(side="right", padx=4)
        self.make_help_button(controls_card, "Enviar-les totes", self.submit_all, "Envia totes les files preparades al formulari.").pack(side="right", padx=4)

        body = ttk.PanedWindow(self.main_tab, orient="horizontal")
        body.pack(fill="both", expand=True, pady=(0, 10))

        left = ttk.Frame(body, style="Card.TFrame", padding=10)
        right_outer = ttk.Frame(body, style="Card.TFrame", padding=0)
        body.add(left, weight=3)
        body.add(right_outer, weight=2)

        filter_bar = ttk.Frame(left, style="Card.TFrame")
        filter_bar.pack(fill="x", pady=(0, 8))
        ttk.Label(filter_bar, text="Registres carregats", style="Section.TLabel").pack(side="left")
        ttk.Label(filter_bar, text="Filtre", style="Section.TLabel").pack(side="right", padx=(8, 4))
        filter_combo = ttk.Combobox(filter_bar, textvariable=self.status_filter_var, values=list(FILTER_OPTIONS), width=18, state="readonly")
        filter_combo.pack(side="right")
        filter_combo.bind("<<ComboboxSelected>>", lambda _event: self.refresh_tree())

        mass_edit_bar = ttk.Frame(left, style="Card.TFrame")
        mass_edit_bar.pack(fill="x", pady=(0, 8))
        ttk.Label(mass_edit_bar, text="EdiciÃ³ massiva", style="Section.TLabel").pack(side="left", padx=(0, 8))
        mass_field_combo = ttk.Combobox(
            mass_edit_bar,
            textvariable=self.mass_edit_field_var,
            values=list(MASS_EDIT_FIELDS.keys()),
            width=24,
            state="readonly",
        )
        mass_field_combo.pack(side="left", padx=4)
        self.mass_edit_value_combo = ttk.Combobox(
            mass_edit_bar,
            textvariable=self.mass_edit_value_var,
            values=FIELD_OPTIONS[MASS_EDIT_FIELDS[self.mass_edit_field_var.get()]],
            width=38,
            state="readonly",
        )
        self.mass_edit_value_combo.pack(side="left", padx=4)
        mass_field_combo.bind("<<ComboboxSelected>>", self.on_mass_edit_field_change)
        self.make_help_button(
            mass_edit_bar,
            "Aplicar a seleccionades",
            self.apply_mass_edit_to_selection,
            "Canvia este camp en totes les files seleccionades.",
        ).pack(side="left", padx=4)

        columns = ("localitat", "data", "categoria", "nom", "preu", "font", "status")
        self.tree = ttk.Treeview(left, columns=columns, show="headings", height=26, selectmode="extended")
        headings = {
            "localitat": "Localitat",
            "data": "Data",
            "categoria": "Categoria",
            "nom": "Activitat",
            "preu": "Preu",
            "font": "Font",
            "status": "Estat",
        }
        widths = {
            "localitat": 110,
            "data": 100,
            "categoria": 150,
            "nom": 360,
            "preu": 80,
            "font": 260,
            "status": 120,
        }
        for column in columns:
            self.tree.heading(column, text=headings[column])
            self.tree.column(column, width=widths[column], anchor="w")
        self.tree.tag_configure("sent", background=COLORS["sent_bg"])
        self.tree.tag_configure("error", background=COLORS["error_bg"])
        self.tree.tag_configure("invalid", background=COLORS["warning_bg"])
        self.tree.pack(side="left", fill="both", expand=True)
        self.tree.bind("<<TreeviewSelect>>", self.on_tree_select)
        self.tree.bind("<Control-a>", self.select_all_visible_rows)
        self.tree.bind("<Control-A>", self.select_all_visible_rows)

        scrollbar = ttk.Scrollbar(left, orient="vertical", command=self.tree.yview)
        scrollbar.pack(side="right", fill="y")
        self.tree.configure(yscrollcommand=scrollbar.set)

        right_canvas = tk.Canvas(right_outer, bg=COLORS["paper"], highlightthickness=0)
        right_scrollbar = ttk.Scrollbar(right_outer, orient="vertical", command=right_canvas.yview)
        right_canvas.configure(yscrollcommand=right_scrollbar.set)
        right_canvas.pack(side="left", fill="both", expand=True)
        right_scrollbar.pack(side="right", fill="y")

        right = ttk.Frame(right_canvas, style="Card.TFrame", padding=16)
        right_window = right_canvas.create_window((0, 0), window=right, anchor="nw")

        def resize_editor_canvas(_event=None):
            right_canvas.configure(scrollregion=right_canvas.bbox("all"))
            right_canvas.itemconfigure(right_window, width=right_canvas.winfo_width())

        def scroll_editor(event):
            right_canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

        right.bind("<Configure>", resize_editor_canvas)
        right_canvas.bind("<Configure>", resize_editor_canvas)
        right_canvas.bind("<Enter>", lambda _event: right_canvas.bind_all("<MouseWheel>", scroll_editor))
        right_canvas.bind("<Leave>", lambda _event: right_canvas.unbind_all("<MouseWheel>"))

        ttk.Label(right, text="Editor de fila", style="Section.TLabel").grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 8))
        validation_label = tk.Label(
            right,
            textvariable=self.validation_var,
            bg=COLORS["warning_bg"],
            fg=COLORS["charcoal"],
            justify="left",
            anchor="w",
            wraplength=360,
            padx=10,
            pady=8,
        )
        validation_label.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(0, 10))

        row_index = 2
        for key, label in FIELD_LABELS.items():
            ttk.Label(right, text=label, style="Body.TLabel").grid(row=row_index, column=0, sticky="nw", pady=5)
            widget = self.build_editor_widget(right, key)
            widget.grid(row=row_index, column=1, sticky="ew", pady=5)
            self.editor_widgets[key] = widget
            row_index += 1

        ttk.Label(right, text="Historial d'enviaments", style="Section.TLabel").grid(row=row_index, column=0, columnspan=2, sticky="w", pady=(12, 6))
        row_index += 1
        self.history_box = tk.Text(
            right,
            height=10,
            wrap="word",
            font=("Consolas", 10),
            bg="white",
            fg=COLORS["charcoal"],
            relief="flat",
            padx=10,
            pady=10,
        )
        self.history_box.grid(row=row_index, column=0, columnspan=2, sticky="nsew")
        self.history_box.configure(state="disabled")

        right.columnconfigure(1, weight=1)
        right.rowconfigure(row_index, weight=1)

        self.build_ocr_tab()

        status = ttk.Label(self.root, textvariable=self.status_var, style="Status.TLabel", padding=(16, 0, 16, 12))
        status.pack(fill="x")

    def build_editor_widget(self, parent, key):
        if key in FIELD_OPTIONS:
            return ttk.Combobox(parent, textvariable=self.editor_vars[key], values=FIELD_OPTIONS[key], state="readonly", width=34)
        width = 42 if key in {"nom", "companyia", "lloc", "font"} else 34
        return ttk.Entry(parent, textvariable=self.editor_vars[key], width=width)

    def make_help_button(self, parent, text, command, help_text):
        button = ttk.Button(parent, text=text, command=command, style="Neutral.TButton")
        button.bind("<Button-3>", lambda event, msg=help_text: self.show_button_help(event, msg))
        return button

    def build_ocr_tab(self):
        outer = ttk.Frame(self.ocr_tab, style="Root.TFrame", padding=(0, 0, 0, 10))
        outer.pack(fill="both", expand=True)

        hero_card = ttk.Frame(outer, style="Card.TFrame", padding=16)
        hero_card.pack(fill="x", pady=(0, 10))
        ttk.Label(hero_card, text="OCR de PDFs", style="Section.TLabel").pack(anchor="w")
        ttk.Label(
            hero_card,
            text="Processa un PDF amb OCR, revisa el text per pàgines i exporta una versió neta a DOCX o PDF.",
            style="Body.TLabel",
            wraplength=1100,
            justify="left",
        ).pack(anchor="w", pady=(6, 0))

        top_card = ttk.Frame(outer, style="Card.TFrame", padding=16)
        top_card.pack(fill="x")
        ttk.Label(top_card, text="PDF origen", style="Section.TLabel").grid(row=0, column=0, sticky="w")
        ttk.Entry(top_card, textvariable=self.ocr_pdf_path_var, width=92).grid(row=0, column=1, sticky="ew", padx=8, pady=4)
        ttk.Button(top_card, text="Buscar PDF", command=self.pick_ocr_pdf, style="Neutral.TButton").grid(row=0, column=2, padx=4)
        ttk.Button(top_card, text="Obrir OCR en PDF24", command=self.process_ocr_pdf, style="Neutral.TButton").grid(row=0, column=3, padx=4)
        ttk.Button(top_card, text="Cancel·lar / netejar", command=self.cancel_ocr_process, style="Neutral.TButton").grid(row=0, column=4, padx=4)
        ttk.Label(top_card, text="PDF OCR", style="Section.TLabel").grid(row=1, column=0, sticky="w")
        ttk.Entry(top_card, textvariable=self.ocr_result_path_var, width=92, state="readonly").grid(row=1, column=1, sticky="ew", padx=8, pady=4)
        ttk.Button(top_card, text="Obrir PDF OCR", command=self.open_ocr_pdf, style="Neutral.TButton").grid(row=1, column=2, padx=4)
        ttk.Button(top_card, text="Carregar PDF OCR", command=self.pick_generated_ocr_pdf, style="Neutral.TButton").grid(row=1, column=3, padx=4)
        ttk.Button(top_card, text="Reextraure text", command=self.reload_text_from_ocr, style="Neutral.TButton").grid(row=1, column=4, padx=4)
        ttk.Label(top_card, text="Localitat", style="Section.TLabel").grid(row=2, column=0, sticky="w")
        ttk.Entry(top_card, textvariable=self.ocr_localitat_var, width=28).grid(row=2, column=1, sticky="w", padx=8, pady=4)
        ttk.Label(top_card, text="Font", style="Section.TLabel").grid(row=2, column=1, sticky="e")
        ttk.Entry(top_card, textvariable=self.ocr_font_var, width=42).grid(row=2, column=2, columnspan=3, sticky="ew", padx=8, pady=4)
        top_card.columnconfigure(1, weight=1)

        summary_card = ttk.Frame(outer, style="Card.TFrame", padding=16)
        summary_card.pack(fill="x", pady=(10, 0))
        ttk.Label(summary_card, text="Resum", style="Section.TLabel").pack(anchor="w")
        ttk.Label(summary_card, textvariable=self.ocr_summary_var, style="Body.TLabel", wraplength=1100, justify="left").pack(anchor="w", pady=(6, 10))
        self.ocr_progress = ttk.Progressbar(summary_card, maximum=100, variable=self.ocr_progress_var, style="Accent.Horizontal.TProgressbar")
        self.ocr_progress.pack(fill="x")

        controls = ttk.Frame(outer, style="Card.TFrame", padding=10)
        controls.pack(fill="x", pady=(10, 0))
        ttk.Button(controls, text="Aplicar canvis de pàgina", command=self.apply_current_ocr_page, style="Neutral.TButton").pack(side="left", padx=4)
        ttk.Button(controls, text="Generar Excel amb ChatGPT", command=self.generate_excel_from_ocr_with_chatgpt, style="Neutral.TButton").pack(side="left", padx=4)
        ttk.Button(controls, text="Configurar IA", command=self.open_ai_settings, style="Neutral.TButton").pack(side="left", padx=4)
        ttk.Button(controls, text="Exportar a PDF net", command=self.export_ocr_to_pdf, style="Neutral.TButton").pack(side="right", padx=4)
        ttk.Button(controls, text="Exportar a DOCX", command=self.export_ocr_to_docx, style="Neutral.TButton").pack(side="right", padx=4)

        body = ttk.PanedWindow(outer, orient="horizontal")
        body.pack(fill="both", expand=True, pady=(10, 0))
        left = ttk.Frame(body, style="Card.TFrame", padding=10)
        center = ttk.Frame(body, style="Card.TFrame", padding=10)
        right = ttk.Frame(body, style="Card.TFrame", padding=10)
        body.add(left, weight=2)
        body.add(center, weight=5)
        body.add(right, weight=2)

        ttk.Label(left, text="Pàgines detectades", style="Section.TLabel").pack(anchor="w", pady=(0, 8))
        self.ocr_page_tree = ttk.Treeview(left, columns=("page", "preview"), show="headings", height=18)
        self.ocr_page_tree.heading("page", text="Pàgina")
        self.ocr_page_tree.heading("preview", text="Vista prèvia")
        self.ocr_page_tree.column("page", width=70, anchor="w")
        self.ocr_page_tree.column("preview", width=240, anchor="w")
        self.ocr_page_tree.pack(side="left", fill="both", expand=True)
        self.ocr_page_tree.bind("<<TreeviewSelect>>", self.on_ocr_page_select)
        page_scroll = ttk.Scrollbar(left, orient="vertical", command=self.ocr_page_tree.yview)
        page_scroll.pack(side="right", fill="y")
        self.ocr_page_tree.configure(yscrollcommand=page_scroll.set)

        ttk.Label(center, text="Editor de text", style="Section.TLabel").pack(anchor="w", pady=(0, 8))
        self.ocr_editor = tk.Text(
            center,
            wrap="word",
            undo=True,
            font=("Consolas", 11),
            bg="white",
            fg=COLORS["charcoal"],
            insertbackground=COLORS["red"],
            relief="flat",
            padx=12,
            pady=12,
        )
        self.ocr_editor.pack(side="left", fill="both", expand=True)
        self.ocr_editor.bind("<<Modified>>", self.on_ocr_editor_modified)
        editor_scroll = ttk.Scrollbar(center, orient="vertical", command=self.ocr_editor.yview)
        editor_scroll.pack(side="right", fill="y")
        self.ocr_editor.configure(yscrollcommand=editor_scroll.set)

        ttk.Label(right, text="Estat del procés", style="Section.TLabel").pack(anchor="w", pady=(0, 8))
        self.ocr_log_box = tk.Text(
            right,
            wrap="word",
            font=("Segoe UI", 10),
            bg="white",
            fg=COLORS["charcoal"],
            relief="flat",
            height=18,
            padx=10,
            pady=10,
        )
        self.ocr_log_box.pack(fill="both", expand=True)
        self.ocr_log_box.configure(state="disabled")

        ttk.Label(outer, textvariable=self.ocr_status_var, style="Status.TLabel", padding=(0, 8, 0, 0)).pack(fill="x")

    def bind_editor_events(self):
        for variable in self.editor_vars.values():
            variable.trace_add("write", self.on_editor_change)

    def bind_shortcuts(self):
        self.root.bind("<Control-o>", lambda _event: self.pick_excel())
        self.root.bind("<Control-l>", lambda _event: self.load_excel())
        self.root.bind("<Control-f>", lambda _event: self.load_form())
        self.root.bind("<Control-s>", lambda _event: self.apply_current_record())
        self.root.bind("<Delete>", lambda _event: self.delete_current_record())
        self.root.bind("<Control-Return>", lambda _event: self.submit_selected())
        self.root.bind("<Control-Shift-Return>", lambda _event: self.submit_all())
        self.root.bind("<Control-g>", lambda _event: self.open_google_search())
        self.root.bind("<Control-u>", lambda _event: self.open_source())
        self.root.bind("<Control-d>", lambda _event: self.apply_fallback_font_to_all())
        self.root.bind("<Control-e>", lambda _event: self.focus_editor())
        self.root.bind("<Control-t>", lambda _event: self.tree.focus_set())
        self.root.bind("<Control-Up>", lambda _event: self.move_selection(-1))
        self.root.bind("<Control-Down>", lambda _event: self.move_selection(1))

    def show_splash(self):
        logo_path = resolve_asset_path(LOGO_FILENAME)
        if not logo_path.exists():
            return
        splash = tk.Toplevel(self.root)
        splash.overrideredirect(True)
        splash.configure(bg=COLORS["cream"])
        splash.attributes("-topmost", True)
        self.root.withdraw()

        width, height = 340, 320
        x = self.root.winfo_screenwidth() // 2 - width // 2
        y = self.root.winfo_screenheight() // 2 - height // 2
        splash.geometry(f"{width}x{height}+{x}+{y}")

        frame = tk.Frame(splash, bg=COLORS["cream"], highlightbackground=COLORS["line"], highlightthickness=1)
        frame.pack(fill="both", expand=True)

        image = Image.open(logo_path).convert("RGBA")
        image.thumbnail((180, 180))
        self.logo_image = ImageTk.PhotoImage(image)

        tk.Label(frame, image=self.logo_image, bg=COLORS["cream"]).pack(pady=(42, 14))
        tk.Label(frame, text=APP_NAME, bg=COLORS["cream"], fg=COLORS["red"], font=("Segoe UI Semibold", 22)).pack()
        tk.Label(frame, text="Iniciant aplicació...", bg=COLORS["cream"], fg=COLORS["muted"], font=("Segoe UI", 10)).pack(pady=(8, 0))

        steps = [88, 90, 92, 94, 96, 98, 100]

        def animate(i=0):
            if i >= len(steps):
                splash.destroy()
                self.root.deiconify()
                self.root.lift()
                try:
                    self.root.attributes("-topmost", True)
                    self.safe_after(120, lambda: self.root.attributes("-topmost", False))
                except tk.TclError:
                    pass
                self.root.focus_force()
                return
            try:
                splash.attributes("-alpha", steps[i] / 100)
            except tk.TclError:
                splash.destroy()
                self.root.deiconify()
                return
            self.safe_after(55, lambda: animate(i + 1))

        self.safe_after(150, animate)

    def now_stamp(self):
        return dt.datetime.now().strftime("%d/%m/%Y %H:%M:%S")

    def append_history(self, text):
        line = f"[{self.now_stamp()}] {text}"
        self.history_entries.append(line)
        self.history_box.configure(state="normal")
        self.history_box.insert("end", line + "\n")
        self.history_box.see("end")
        self.history_box.configure(state="disabled")

    def set_status(self, text):
        self.status_var.set(text)

    def show_button_help(self, event, text):
        self.hide_button_help()
        popup = tk.Toplevel(self.root)
        popup.overrideredirect(True)
        popup.attributes("-topmost", True)
        popup.configure(bg=COLORS["charcoal"])
        label = tk.Label(
            popup,
            text=text,
            bg=COLORS["charcoal"],
            fg="white",
            font=("Segoe UI", 9),
            padx=10,
            pady=6,
        )
        label.pack()
        popup.geometry(f"+{event.x_root + 8}+{event.y_root + 8}")
        self.help_popup = popup
        self.root.after(2200, self.hide_button_help)
        return "break"

    def hide_button_help(self):
        if self.help_popup is not None:
            try:
                self.help_popup.destroy()
            except tk.TclError:
                pass
            self.help_popup = None

    def safe_after(self, delay_ms, callback):
        if self.closing:
            return None
        after_id = None

        def run_callback():
            self.after_ids.discard(after_id)
            if self.closing:
                return
            try:
                if self.root.winfo_exists():
                    callback()
            except tk.TclError:
                pass

        try:
            after_id = self.root.after(delay_ms, run_callback)
        except tk.TclError:
            return None
        self.after_ids.add(after_id)
        return after_id

    def on_close(self):
        if self.pdf24_process is not None and self.pdf24_process.poll() is None:
            try:
                self.pdf24_process.terminate()
            except Exception:
                pass
        self.closing = True
        for after_id in list(self.after_ids):
            try:
                self.root.after_cancel(after_id)
            except tk.TclError:
                pass
        self.after_ids.clear()
        try:
            if self.autosave_after_id:
                self.root.after_cancel(self.autosave_after_id)
        except tk.TclError:
            pass
        self.root.destroy()

    def append_ocr_log(self, text):
        def _write():
            self.ocr_log_box.configure(state="normal")
            self.ocr_log_box.insert("end", f"{text}\n")
            self.ocr_log_box.see("end")
            self.ocr_log_box.configure(state="disabled")
            self.ocr_status_var.set(text)
            percent_match = re.search(r"(?<!\d)(\d{1,3})(?:[.,]\d+)?\s*%", text)
            if percent_match:
                percent = max(0, min(100, int(percent_match.group(1))))
                self.ocr_progress_var.set(percent)

        self.safe_after(0, _write)

    def clear_ocr_log(self):
        self.ocr_log_box.configure(state="normal")
        self.ocr_log_box.delete("1.0", "end")
        self.ocr_log_box.configure(state="disabled")

    def reset_ocr_state(self, message="Procés OCR netejat. Selecciona un PDF per començar."):
        self.source_pdf = None
        self.ocr_pdf = None
        self.ocr_pages = []
        self.ocr_current_page_index = None
        self.ocr_dirty = False
        self.ocr_pdf_path_var.set("")
        self.ocr_result_path_var.set("Encara no s'ha generat.")
        self.ocr_summary_var.set(message)
        self.ocr_status_var.set(message)
        self.ocr_progress_var.set(0)
        for item in self.ocr_page_tree.get_children():
            self.ocr_page_tree.delete(item)
        self.ocr_editor.delete("1.0", "end")
        self.clear_ocr_log()
        self.append_ocr_log(message)

    def cancel_ocr_process(self):
        if self.pdf24_process is not None and self.pdf24_process.poll() is None:
            try:
                self.pdf24_process.terminate()
            except Exception:
                pass
            self.pdf24_process = None
        self.reset_ocr_state()

    def pick_ocr_pdf(self):
        path = filedialog.askopenfilename(
            title="Selecciona un PDF",
            filetypes=[("PDF", "*.pdf"), ("Tots", "*.*")],
        )
        if path:
            self.source_pdf = Path(path)
            self.ocr_pdf_path_var.set(path)
            size_mb = self.source_pdf.stat().st_size / (1024 * 1024)
            self.ocr_summary_var.set(f"PDF carregat: {self.source_pdf.name} | {size_mb:.2f} MB")
            self.ocr_progress_var.set(0)
            self.append_ocr_log("PDF seleccionat.")

    def open_ocr_pdf(self):
        if not self.ocr_pdf or not self.ocr_pdf.exists():
            messagebox.showinfo(APP_NAME, "Encara no hi ha cap PDF OCR generat.")
            return
        os.startfile(self.ocr_pdf)  # type: ignore[attr-defined]

    def pick_generated_ocr_pdf(self):
        path = filedialog.askopenfilename(
            title="Selecciona el PDF OCR guardat des de PDF24",
            filetypes=[("PDF", "*.pdf"), ("Tots", "*.*")],
        )
        if not path:
            return
        self.ocr_pdf = Path(path)
        self.ocr_result_path_var.set(path)
        self.reload_text_from_ocr()

    def load_ocr_pages_into_ui(self):
        for item in self.ocr_page_tree.get_children():
            self.ocr_page_tree.delete(item)
        for page in self.ocr_pages:
            preview = page.text.splitlines()[0] if page.text.strip() else "(sense text)"
            self.ocr_page_tree.insert("", "end", iid=str(page.number - 1), values=(page.number, preview[:48]))
        if self.ocr_pages:
            self.ocr_page_tree.selection_set("0")
            self.ocr_page_tree.focus("0")
            self.populate_ocr_editor(0)

    def populate_ocr_editor(self, index: int):
        self.ocr_current_page_index = index
        self.ocr_suspend_dirty = True
        try:
            self.ocr_editor.delete("1.0", "end")
            self.ocr_editor.insert("1.0", self.ocr_pages[index].text)
            self.ocr_editor.edit_modified(False)
        finally:
            self.ocr_suspend_dirty = False
        self.ocr_dirty = False

    def on_ocr_page_select(self, _event):
        selection = self.ocr_page_tree.selection()
        if not selection:
            return
        target = int(selection[0])
        if self.ocr_current_page_index is not None and target != self.ocr_current_page_index and self.ocr_dirty:
            self.apply_current_ocr_page()
        self.populate_ocr_editor(target)

    def on_ocr_editor_modified(self, _event):
        if self.ocr_suspend_dirty:
            self.ocr_editor.edit_modified(False)
            return
        self.ocr_dirty = True
        self.ocr_status_var.set("Hi ha canvis pendents d'aplicar en la pàgina actual.")
        self.ocr_editor.edit_modified(False)

    def apply_current_ocr_page(self):
        if self.ocr_current_page_index is None or not self.ocr_pages:
            return True
        self.ocr_pages[self.ocr_current_page_index].text = self.ocr_editor.get("1.0", "end-1c")
        preview = self.ocr_pages[self.ocr_current_page_index].text.splitlines()[0] if self.ocr_pages[self.ocr_current_page_index].text.strip() else "(sense text)"
        self.ocr_page_tree.item(
            str(self.ocr_current_page_index),
            values=(self.ocr_pages[self.ocr_current_page_index].number, preview[:48]),
        )
        self.ocr_dirty = False
        self.ocr_status_var.set("Canvis aplicats.")
        return True

    def process_ocr_pdf(self):
        source_text = self.ocr_pdf_path_var.get().strip()
        if source_text:
            self.source_pdf = Path(source_text)
        if not self.source_pdf or not self.source_pdf.exists():
            messagebox.showwarning(APP_NAME, "Selecciona abans un PDF.")
            return
        pdf24_ocr_exe = resolve_pdf24_ocr_exe()
        if not pdf24_ocr_exe.exists():
            messagebox.showerror(
                APP_NAME,
                "No he trobat PDF24 OCR.\n\n"
                "Instal·la PDF24 Creator i torna-ho a provar.\n"
                "L'instal·lador d'AutoCPV també intenta instal·lar-lo amb winget si falta.",
            )
            return

        self.ocr_pdf = None
        self.ocr_pages = []
        self.ocr_current_page_index = None
        self.ocr_dirty = False
        self.ocr_result_path_var.set("Guarda el resultat des de PDF24 i després prem 'Carregar PDF OCR'.")
        self.ocr_progress_var.set(0)
        for item in self.ocr_page_tree.get_children():
            self.ocr_page_tree.delete(item)
        self.ocr_editor.delete("1.0", "end")
        self.clear_ocr_log()

        try:
            self.pdf24_process = subprocess.Popen(
                [str(pdf24_ocr_exe), str(self.source_pdf)],
                cwd=str(pdf24_ocr_exe.parent),
            )
        except Exception as exc:
            messagebox.showerror(APP_NAME, f"No he pogut obrir PDF24 OCR: {exc}")
            return

        self.ocr_summary_var.set("PDF24 OCR està obert amb el PDF seleccionat. Fes l'OCR en PDF24 i guarda el resultat.")
        self.ocr_status_var.set("PDF24 OCR obert. Esperant que carregues el PDF OCR resultant.")
        self.append_ocr_log("PDF24 OCR obert amb el PDF seleccionat.")
        self.append_ocr_log("Quan PDF24 acabe, guarda el PDF OCR i torna ací per prémer 'Carregar PDF OCR'.")
    def reload_text_from_ocr(self):
        if not self.ocr_pdf or not self.ocr_pdf.exists():
            messagebox.showinfo(APP_NAME, "Encara no hi ha cap PDF OCR per a reprocessar.")
            return
        self.apply_current_ocr_page()
        try:
            self.ocr_pages = extract_document_pages(self.ocr_pdf)
            self.load_ocr_pages_into_ui()
            self.ocr_summary_var.set(f"Text reextret: {len(self.ocr_pages)} pàgines.")
            self.append_ocr_log("Text reextret des del PDF OCR.")
        except Exception as exc:
            messagebox.showerror(APP_NAME, str(exc))

    def export_ocr_to_docx(self):
        if not self.ocr_pages:
            messagebox.showwarning(APP_NAME, "No hi ha text carregat per a exportar.")
            return
        self.apply_current_ocr_page()
        output = filedialog.asksaveasfilename(
            title="Guardar DOCX",
            defaultextension=".docx",
            filetypes=[("Word", "*.docx")],
        )
        if not output:
            return
        export_docx(self.ocr_pages, Path(output))
        self.append_ocr_log(f"DOCX exportat: {output}")

    def export_ocr_to_pdf(self):
        if not self.ocr_pages:
            messagebox.showwarning(APP_NAME, "No hi ha text carregat per a exportar.")
            return
        self.apply_current_ocr_page()
        output = filedialog.asksaveasfilename(
            title="Guardar PDF net",
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
        )
        if not output:
            return
        export_clean_pdf(self.ocr_pages, Path(output))
        self.append_ocr_log(f"PDF net exportat: {output}")

    def open_ai_settings(self):
        window = tk.Toplevel(self.root)
        window.title("Configuracio IA")
        window.geometry("620x260")
        window.configure(bg=COLORS["paper"])
        window.transient(self.root)
        window.grab_set()

        frame = ttk.Frame(window, style="Card.TFrame", padding=18)
        frame.pack(fill="both", expand=True)
        ttk.Label(frame, text="NVIDIA Integrate", style="Section.TLabel").grid(row=0, column=0, columnspan=2, sticky="w", pady=(0, 12))

        ttk.Label(frame, text="API key", style="Body.TLabel").grid(row=1, column=0, sticky="w", pady=6)
        key_entry = ttk.Entry(frame, textvariable=self.nvidia_api_key_var, width=64, show="*")
        key_entry.grid(row=1, column=1, sticky="ew", pady=6)

        ttk.Label(frame, text="Model", style="Body.TLabel").grid(row=2, column=0, sticky="w", pady=6)
        ttk.Entry(frame, textvariable=self.nvidia_model_var, width=64).grid(row=2, column=1, sticky="ew", pady=6)

        ttk.Label(frame, text="Max tokens", style="Body.TLabel").grid(row=3, column=0, sticky="w", pady=6)
        ttk.Entry(frame, textvariable=self.nvidia_max_tokens_var, width=18).grid(row=3, column=1, sticky="w", pady=6)

        ttk.Label(frame, text=f"Es guarda en {CONFIG_PATH}", style="Body.TLabel", wraplength=560).grid(row=4, column=0, columnspan=2, sticky="w", pady=(8, 12))

        buttons = ttk.Frame(frame, style="Card.TFrame")
        buttons.grid(row=5, column=0, columnspan=2, sticky="e")

        def save_settings():
            api_key = self.nvidia_api_key_var.get().strip()
            model = self.nvidia_model_var.get().strip() or DEFAULT_NVIDIA_MODEL
            try:
                max_tokens = int(self.nvidia_max_tokens_var.get().strip() or DEFAULT_NVIDIA_MAX_TOKENS)
            except ValueError:
                messagebox.showwarning(APP_NAME, "Max tokens ha de ser un numero.")
                return
            if max_tokens < 1000:
                messagebox.showwarning(APP_NAME, "Max tokens hauria de ser com a minim 1000.")
                return
            self.nvidia_model_var.set(model)
            self.nvidia_max_tokens_var.set(str(max_tokens))
            self.config_data["ai"] = {
                "nvidia_api_key": api_key,
                "nvidia_model": model,
                "nvidia_max_tokens": max_tokens,
            }
            save_app_config(self.config_data)
            self.append_ocr_log("Configuracio IA guardada.")
            window.destroy()

        ttk.Button(buttons, text="Guardar", command=save_settings, style="Neutral.TButton").pack(side="right", padx=4)
        ttk.Button(buttons, text="Cancel·lar", command=window.destroy, style="Neutral.TButton").pack(side="right", padx=4)
        frame.columnconfigure(1, weight=1)
        key_entry.focus_set()

    def generate_excel_from_ocr_with_chatgpt(self):
        if not self.ocr_pages:
            messagebox.showwarning(APP_NAME, "Carrega primer un PDF OCR i extrau el text.")
            return
        localitat = self.ocr_localitat_var.get().strip()
        if not localitat:
            messagebox.showwarning(APP_NAME, "Escriu la localitat abans de generar l'Excel.")
            return
        prompt_path = resolve_prompt_path()
        if not prompt_path.exists():
            messagebox.showerror(APP_NAME, f"No he trobat el prompt en {prompt_path}.")
            return
        api_key = self.nvidia_api_key_var.get().strip() or os.environ.get("NVIDIA_API_KEY", "").strip()
        if not api_key:
            messagebox.showinfo(
                APP_NAME,
                "Encara no hi ha clau de NVIDIA configurada.\n\n"
                "Prem 'Configurar IA' i guarda la clau NVIDIA per a este ordinador.",
            )
            self.append_ocr_log("Falta NVIDIA_API_KEY. La generacio amb ChatGPT queda preparada pero no s'ha executat.")
            return
        model = self.nvidia_model_var.get().strip() or DEFAULT_NVIDIA_MODEL
        try:
            max_tokens = int(self.nvidia_max_tokens_var.get().strip() or DEFAULT_NVIDIA_MAX_TOKENS)
        except ValueError:
            messagebox.showwarning(APP_NAME, "Revisa Configurar IA: Max tokens ha de ser un numero.")
            return

        self.apply_current_ocr_page()
        output = filedialog.asksaveasfilename(
            title="Guardar Excel generat per ChatGPT",
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx")],
            initialfile=f"{localitat}_autocpv.xlsx",
        )
        if not output:
            return

        prompt_text = prompt_path.read_text(encoding="utf-8")
        prompt_text = prompt_text.replace("{{LOCALITAT}}", localitat)
        prompt_text = prompt_text.replace("{{FONT}}", self.ocr_font_var.get().strip())
        prompt_text = prompt_text.replace("{{FITXER_BASE_OPCIONAL}}", "")
        ocr_text = pages_to_prompt_text(self.ocr_pages)

        def worker():
            try:
                self.safe_after(0, lambda: self.ocr_progress_var.set(10))
                self.append_ocr_log("Preparant text OCR per a ChatGPT...")
                result = call_ai_for_autocpv(prompt_text, ocr_text, api_key, self.append_ocr_log, model=model, max_tokens=max_tokens)
                rows = result.get("rows", [])
                if not isinstance(rows, list):
                    raise RuntimeError("La resposta de ChatGPT no conté una llista de files.")
                write_autocpv_excel(rows, Path(output))
                report = result.get("report", {})
                self.safe_after(0, lambda: self.ocr_progress_var.set(100))
                self.append_ocr_log(f"Excel generat: {output}")
                self.append_ocr_log(f"Files generades: {len(rows)}")
                discarded = report.get("discarded", []) if isinstance(report, dict) else []
                if discarded:
                    self.append_ocr_log(f"Descartades rellevants: {len(discarded)}")
                self.safe_after(0, lambda: self.ocr_summary_var.set(f"Excel generat amb ChatGPT: {Path(output).name} | Files: {len(rows)}"))
                self.safe_after(0, lambda: self.excel_path_var.set(output))
                self.safe_after(0, self.load_excel)
            except Exception as exc:
                self.append_ocr_log(f"Error generant Excel amb ChatGPT: {exc}")
                self.safe_after(0, lambda: messagebox.showerror(APP_NAME, str(exc)))
                self.safe_after(0, lambda: self.ocr_progress_var.set(0))

        threading.Thread(target=worker, daemon=True).start()
        self.append_ocr_log("Generació amb ChatGPT iniciada...")

    def show_shortcuts_help(self):
        window = tk.Toplevel(self.root)
        window.title("Tecles ràpides")
        window.geometry("520x500")
        window.configure(bg=COLORS["paper"])

        frame = ttk.Frame(window, style="Card.TFrame", padding=18)
        frame.pack(fill="both", expand=True)
        ttk.Label(frame, text="Tecles ràpides", style="Section.TLabel").pack(anchor="w", pady=(0, 10))

        shortcuts = [
            ("Ctrl + O", "Buscar Excel"),
            ("Ctrl + L", "Carregar Excel"),
            ("Ctrl + F", "Llegir formulari"),
            ("Ctrl + S", "Aplicar canvis de la fila"),
            ("Supr", "Eliminar fila seleccionada"),
            ("Ctrl + Enter", "Enviar fila seleccionada"),
            ("Ctrl + Shift + Enter", "Enviar totes les files"),
            ("Ctrl + G", "Buscar a Google"),
            ("Ctrl + U", "Obrir font"),
            ("Ctrl + D", "Aplicar font per defecte a totes"),
            ("Ctrl + E", "Portar el focus a l'editor"),
            ("Ctrl + T", "Portar el focus a la taula"),
            ("Ctrl + A", "Seleccionar totes les files visibles de la taula"),
            ("Ctrl + Arriba / Baix", "Canviar de fila"),
        ]

        box = tk.Text(
            frame,
            wrap="word",
            font=("Segoe UI", 10),
            bg="white",
            fg=COLORS["charcoal"],
            relief="flat",
            padx=12,
            pady=12,
        )
        box.pack(fill="both", expand=True)
        box.insert("1.0", "\n".join(f"{combo}: {desc}" for combo, desc in shortcuts))
        box.configure(state="disabled")

    def open_progress_window(self, title: str, total: int):
        window = tk.Toplevel(self.root)
        window.title(title)
        window.geometry("460x150")
        window.resizable(False, False)
        window.configure(bg=COLORS["paper"])
        window.transient(self.root)
        window.grab_set()

        outer = tk.Frame(window, bg=COLORS["paper"], padx=18, pady=18)
        outer.pack(fill="both", expand=True)
        tk.Label(outer, text=title, bg=COLORS["paper"], fg=COLORS["red"], font=("Segoe UI Semibold", 13), anchor="w").pack(fill="x")
        status_var = tk.StringVar(value="Preparant procés...")
        tk.Label(outer, textvariable=status_var, bg=COLORS["paper"], fg=COLORS["charcoal"], font=("Segoe UI", 10), anchor="w").pack(fill="x", pady=(10, 8))
        progress_var = tk.DoubleVar(value=0)
        progress = ttk.Progressbar(outer, variable=progress_var, maximum=max(total, 1), style="Accent.Horizontal.TProgressbar", length=400)
        progress.pack(fill="x")
        count_var = tk.StringVar(value=f"0 / {total}")
        tk.Label(outer, textvariable=count_var, bg=COLORS["paper"], fg=COLORS["muted"], font=("Segoe UI", 9), anchor="e").pack(fill="x", pady=(8, 0))
        window.update_idletasks()
        return window, progress_var, status_var, count_var

    def focus_editor(self):
        for key in ("nom", "font", "localitat", "data"):
            widget = self.editor_widgets.get(key)
            if widget is not None:
                widget.focus_set()
                return

    def move_selection(self, delta):
        items = self.tree.get_children()
        if not items:
            return
        selection = self.tree.selection()
        current = selection[0] if selection else items[0]
        try:
            position = items.index(current)
        except ValueError:
            position = 0
        target = items[max(0, min(len(items) - 1, position + delta))]
        self.tree.selection_set(target)
        self.tree.focus(target)
        self.tree.see(target)
        self.on_tree_select(None)

    def select_all_visible_rows(self, _event=None):
        items = self.tree.get_children()
        if not items:
            return "break"
        self.tree.selection_set(items)
        self.tree.focus(items[0])
        self.tree.see(items[0])
        self.populate_editor(int(items[0]))
        self.set_status(f"{len(items)} files seleccionades.")
        return "break"

    def pick_excel(self):
        path = filedialog.askopenfilename(title="Selecciona un Excel", filetypes=[("Excel", "*.xlsx *.xlsm"), ("Tots", "*.*")])
        if path:
            self.excel_path_var.set(path)

    def load_excel(self):
        path = self.excel_path_var.get().strip()
        if not path:
            messagebox.showwarning("Falta l'Excel", "Selecciona un arxiu Excel.")
            return
        try:
            self.records = load_excel_records(path, self.person_var.get(), self.fallback_font_var.get())
        except Exception as exc:
            messagebox.showerror("Error llegint l'Excel", str(exc))
            return
        self.current_index = None
        self.last_deleted = None
        self.refresh_tree()
        self.append_history(f"Excel carregat amb {len(self.records)} files.")
        self.set_status(f"Excel carregat: {len(self.records)} files.")

    def load_form(self):
        form_url = self.form_url_var.get().strip()
        if not form_url:
            messagebox.showwarning("Falta la URL", "Escriu la URL del formulari.")
            return

        def task():
            try:
                metadata = extract_form_metadata(form_url)
            except Exception as exc:
                self.root.after(0, lambda: messagebox.showerror("Error llegint el formulari", str(exc)))
                return
            self.form_metadata = metadata
            self.root.after(0, lambda: self.set_status("Formulari llegit. Camps detectats i preparat per a enviar."))
            self.root.after(0, lambda: self.append_history("Formulari llegit correctament."))

        threading.Thread(target=task, daemon=True).start()
        self.set_status("Llegint formulari...")

    def save_session(self):
        self.apply_current_record(silent=True)
        path = filedialog.asksaveasfilename(
            title="Guardar sessió",
            defaultextension=".autocpv.json",
            filetypes=SESSION_FILETYPES,
        )
        if not path:
            return
        payload = {
            "app": APP_NAME,
            "version": APP_VERSION,
            "saved_at": self.now_stamp(),
            "excel_path": self.excel_path_var.get().strip(),
            "form_url": self.form_url_var.get().strip(),
            "person": self.person_var.get().strip(),
            "fallback_font": self.fallback_font_var.get().strip(),
            "status_filter": self.status_filter_var.get().strip(),
            "records": [asdict(record) for record in self.records],
            "history": self.history_entries,
        }
        Path(path).write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        self.append_history(f"Sessió guardada en {path}.")
        self.set_status("Sessió guardada correctament.")

    def load_session(self):
        path = filedialog.askopenfilename(title="Obrir sessió", filetypes=SESSION_FILETYPES)
        if not path:
            return
        data = json.loads(Path(path).read_text(encoding="utf-8"))
        self.records = [Record(**item) for item in data.get("records", [])]
        self.excel_path_var.set(data.get("excel_path", ""))
        self.form_url_var.set(data.get("form_url", DEFAULT_FORM_URL))
        self.person_var.set(data.get("person", "Pol"))
        self.fallback_font_var.set(data.get("fallback_font", ""))
        self.status_filter_var.set(data.get("status_filter", "Totes") or "Totes")
        self.history_entries = list(data.get("history", []))
        self.reload_history_box()
        self.current_index = None
        self.last_deleted = None
        self.refresh_tree()
        self.append_history(f"Sessió reoberta des de {path}.")
        self.set_status("Sessió carregada correctament.")

    def reload_history_box(self):
        self.history_box.configure(state="normal")
        self.history_box.delete("1.0", "end")
        if self.history_entries:
            self.history_box.insert("end", "\n".join(self.history_entries) + "\n")
        self.history_box.configure(state="disabled")

    def apply_fallback_font_to_all(self):
        self.apply_fallback_font(mode="all")

    def apply_fallback_font_to_empty(self):
        self.apply_fallback_font(mode="empty")

    def on_mass_edit_field_change(self, _event=None):
        field_key = MASS_EDIT_FIELDS.get(self.mass_edit_field_var.get(), "llengua")
        values = FIELD_OPTIONS.get(field_key, [])
        self.mass_edit_value_combo.configure(values=values)
        self.mass_edit_value_var.set(values[0] if values else "")

    def selected_record_indexes(self):
        indexes = []
        for item in self.tree.selection():
            try:
                indexes.append(int(item))
            except ValueError:
                continue
        return indexes

    def apply_mass_edit_to_selection(self):
        selected_indexes = self.selected_record_indexes()
        if not selected_indexes:
            messagebox.showinfo("Sense selecció", "Selecciona una o més files de la taula.")
            return

        field_key = MASS_EDIT_FIELDS.get(self.mass_edit_field_var.get())
        value = self.mass_edit_value_var.get().strip()
        if not field_key or not value:
            messagebox.showwarning("Falta valor", "Tria un camp i un valor abans d'aplicar el canvi.")
            return

        if self.dirty:
            self.apply_current_record(silent=True)

        changed = 0
        for index in selected_indexes:
            if index < 0 or index >= len(self.records):
                continue
            record = self.records[index]
            if getattr(record, field_key) == value:
                continue
            setattr(record, field_key, value)
            record.status = "Pendent"
            record.status_detail = ""
            changed += 1

        if self.current_index in selected_indexes:
            self.suspend_dirty = True
            try:
                self.editor_vars[field_key].set(getattr(self.records[self.current_index], field_key))
            finally:
                self.suspend_dirty = False
            self.dirty = False

        self.refresh_tree(preserve_selection=True)
        self.update_validation_message()
        label = self.mass_edit_field_var.get()
        self.set_status(f"{label} canviada en {changed} files seleccionades.")
        self.append_history(f"Edició massiva: {label} = {value} en {changed} files.")

    def apply_fallback_font(self, mode: str):
        fallback_font = self.fallback_font_var.get().strip()
        if not fallback_font:
            messagebox.showwarning("Falta la font", "Escriu una font per defecte abans d'aplicar-la.")
            return
        if not self.records:
            messagebox.showwarning("Sense files", "Carrega primer un Excel.")
            return

        changed = 0
        for record in self.records:
            if mode == "all" or not record.font.strip():
                record.font = fallback_font
                changed += 1

        if self.current_index is not None and changed:
            self.suspend_dirty = True
            try:
                self.editor_vars["font"].set(self.records[self.current_index].font)
            finally:
                self.suspend_dirty = False
            self.dirty = False

        self.refresh_tree()
        if mode == "all":
            self.set_status("Font per defecte aplicada a totes les files.")
            self.append_history("Font per defecte aplicada a totes les files.")
        else:
            self.set_status(f"Font per defecte aplicada a {changed} files buides.")
            self.append_history(f"Font per defecte aplicada a {changed} files buides.")

    def on_editor_change(self, *_args):
        if self.suspend_dirty or self.current_index is None:
            return
        self.dirty = True
        self.set_status("Guardant canvis automàticament...")
        if self.autosave_after_id is not None:
            self.root.after_cancel(self.autosave_after_id)
        self.autosave_after_id = self.root.after(350, self.autosave_current_record)

    def autosave_current_record(self):
        self.autosave_after_id = None
        self.apply_current_record(silent=True)
        self.set_status("Canvis guardats automàticament.")

    def refresh_tree(self, preserve_selection=True):
        selected_items = []
        if preserve_selection:
            selected_items = list(self.tree.selection())
            if not selected_items and self.current_index is not None:
                selected_items = [str(self.current_index)]

        for item in self.tree.get_children():
            self.tree.delete(item)

        self.visible_indices = [idx for idx, record in enumerate(self.records) if self.record_matches_filter(record)]
        for idx in self.visible_indices:
            record = self.records[idx]
            self.tree.insert(
                "",
                "end",
                iid=str(idx),
                values=(record.localitat, record.data, record.categoria, record.nom, record.preu, record.font, self.display_status(record)),
                tags=(self.record_tag(record),),
            )

        if not self.visible_indices:
            self.clear_editor()
            self.validation_var.set("No hi ha files visibles amb el filtre actual.")
            return

        available_items = set(self.tree.get_children())
        selected_items = [item for item in selected_items if item in available_items]
        if not selected_items:
            selected_items = [str(self.visible_indices[0])]
        self.tree.selection_set(selected_items)
        focused = selected_items[0]
        self.tree.focus(focused)
        self.tree.see(focused)
        self.populate_editor(int(focused))

    def record_matches_filter(self, record: Record) -> bool:
        mode = FILTER_OPTIONS.get(self.status_filter_var.get(), "all")
        tag = self.record_tag(record)
        if mode == "all":
            return True
        if mode == "sent":
            return tag == "sent"
        if mode == "error":
            return tag == "error"
        if mode == "invalid":
            return tag == "invalid"
        return tag not in {"sent", "error", "invalid"}

    def display_status(self, record: Record) -> str:
        if record.status == "Enviat":
            return "Enviat"
        if record.status == "Error":
            return "Error"
        errors = self.validate_record(record)
        if errors:
            return "Revisar"
        return "Pendent"

    def record_tag(self, record: Record) -> str:
        if record.status == "Enviat":
            return "sent"
        if record.status == "Error":
            return "error"
        if self.validate_record(record):
            return "invalid"
        return ""

    def populate_editor(self, index):
        self.current_index = index
        record = self.records[index]
        self.suspend_dirty = True
        try:
            for key in FIELD_LABELS:
                self.editor_vars[key].set(getattr(record, key))
        finally:
            self.suspend_dirty = False
        self.dirty = False
        self.update_validation_summary(index)

    def clear_editor(self):
        self.current_index = None
        self.suspend_dirty = True
        try:
            for key in FIELD_LABELS:
                self.editor_vars[key].set("")
            self.editor_vars["persona"].set(self.person_var.get().strip() or "Pol")
        finally:
            self.suspend_dirty = False
        self.dirty = False

    def apply_current_record(self, silent=False):
        if self.current_index is None:
            return True
        record = self.records[self.current_index]
        previous_values = {key: getattr(record, key) for key in FIELD_LABELS}
        for key in FIELD_LABELS:
            setattr(record, key, self.editor_vars[key].get().strip())
        changed = any(previous_values[key] != getattr(record, key) for key in FIELD_LABELS)
        if changed:
            record.status = "Pendent"
            record.status_detail = ""
        self.update_tree_item(self.current_index)
        self.update_validation_summary(self.current_index)
        self.dirty = False
        if not silent:
            self.set_status("Canvis aplicats correctament.")
        return True

    def update_tree_item(self, index):
        iid = str(index)
        if iid not in self.tree.get_children():
            return
        record = self.records[index]
        self.tree.item(
            iid,
            values=(record.localitat, record.data, record.categoria, record.nom, record.preu, record.font, self.display_status(record)),
            tags=(self.record_tag(record),),
        )

    def update_validation_summary(self, index):
        record = self.records[index]
        errors = self.validate_record(record)
        if errors:
            self.validation_var.set("Cal revisar:\n- " + "\n- ".join(errors[:6]))
        else:
            detail = record.status_detail or "Fila preparada per a enviar."
            self.validation_var.set(detail)

    def validate_record(self, record: Record):
        errors = []
        for key in REQUIRED_FIELDS:
            if not getattr(record, key).strip():
                errors.append(f"Falta {FIELD_LABELS[key].lower()}.")
        if record.categoria == "Altres" and not record.altres.strip():
            errors.append("Has omplit 'Categoria' amb 'Altres' però falta la descripció.")
        if record.data and not is_valid_date(record.data):
            errors.append("La data ha d'estar en format YYYY-MM-DD.")
        if record.preu and not is_valid_price(record.preu):
            errors.append("El preu ha de ser numèric sense símbols.")
        if record.font and not re.match(r"^https?://", record.font.strip()):
            errors.append("La font ha de començar per http:// o https://.")
        return errors

    def delete_current_record(self):
        selected_indexes = self.selected_record_indexes()
        if not selected_indexes and self.current_index is not None:
            selected_indexes = [self.current_index]
        selected_indexes = sorted(set(index for index in selected_indexes if 0 <= index < len(self.records)))
        if not selected_indexes or not self.records:
            messagebox.showinfo("Sense selecció", "Selecciona una o més files per a eliminar.")
            return
        self.apply_current_record(silent=True)
        if len(selected_indexes) == 1:
            record = self.records[selected_indexes[0]]
            description = record.nom or record.localitat or f"fila {selected_indexes[0] + 1}"
            message = f"Vols eliminar completament esta fila?\n\n{description}"
        else:
            preview = []
            for index in selected_indexes[:5]:
                record = self.records[index]
                preview.append(f"- {record.nom or record.localitat or f'fila {index + 1}'}")
            if len(selected_indexes) > 5:
                preview.append(f"- ... i {len(selected_indexes) - 5} més")
            description = f"{len(selected_indexes)} files"
            message = "Vols eliminar completament estes files?\n\n" + "\n".join(preview)
        confirmed = messagebox.askyesno("Eliminar fila", message)
        if not confirmed:
            return

        deleted_records = []
        for index in sorted(selected_indexes, reverse=True):
            deleted_records.append((index, self.records.pop(index)))
        deleted_records.reverse()
        self.last_deleted = deleted_records
        self.current_index = None
        self.refresh_tree()
        self.append_history(f"Eliminació: {description}.")
        if self.records:
            self.set_status(f"{len(deleted_records)} files eliminades correctament. Pots desfer-ho.")
        else:
            self.set_status("Files eliminades. Ja no queden registres carregats.")

    def undo_delete_record(self):
        if not self.last_deleted:
            messagebox.showinfo("Sense canvis", "No hi ha cap eliminació recent per a desfer.")
            return
        deleted_records = self.last_deleted
        if isinstance(deleted_records, tuple):
            deleted_records = [deleted_records]
        restored_indexes = []
        for index, record in sorted(deleted_records, key=lambda item: item[0]):
            index = max(0, min(index, len(self.records)))
            self.records.insert(index, record)
            restored_indexes.append(index)
        self.last_deleted = None
        self.refresh_tree()
        restored_items = [str(index) for index in restored_indexes if str(index) in self.tree.get_children()]
        if restored_items:
            self.tree.selection_set(restored_items)
            self.tree.focus(restored_items[0])
            self.tree.see(restored_items[0])
            self.populate_editor(int(restored_items[0]))
        self.append_history(f"Eliminació desfeta: {len(restored_indexes)} files recuperades.")
        self.set_status(f"{len(restored_indexes)} files recuperades correctament.")

    def on_tree_select(self, _event):
        selection = self.tree.selection()
        if not selection:
            return
        target = int(selection[0])
        if self.current_index is not None and target != self.current_index and self.dirty:
            self.apply_current_record(silent=True)
        self.populate_editor(target)

    def build_search_query(self, record: Record):
        return " ".join(part for part in [record.nom, record.companyia, record.localitat] if part).strip()

    def build_google_lucky_url(self, query: str):
        return "https://www.google.com/search?btnI=I&q=" + urllib.parse.quote(query)

    def resolve_google_first_result(self, query: str):
        if not query.strip():
            return ""
        request = urllib.request.Request(
            self.build_google_lucky_url(query),
            headers={
                "User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/123.0 Safari/537.36"
                )
            },
        )
        with urllib.request.urlopen(request, timeout=20) as response:
            return response.geturl()

    def open_google_search(self):
        if self.current_index is None:
            return
        self.apply_current_record(silent=True)
        record = self.records[self.current_index]
        query = self.build_search_query(record)
        if query:
            webbrowser.open("https://www.google.com/search?q=" + urllib.parse.quote(query))

    def apply_google_first_result_to_selected(self):
        if self.current_index is None:
            messagebox.showinfo("Sense selecció", "Selecciona una fila.")
            return
        self.apply_current_record(silent=True)
        record = self.records[self.current_index]
        query = self.build_search_query(record)
        if not query:
            messagebox.showwarning("Sense dades", "Esta fila no té prou dades per buscar una font.")
            return
        try:
            source_url = self.resolve_google_first_result(query)
        except Exception as exc:
            messagebox.showerror("Error buscant la font", str(exc))
            return
        if not source_url:
            messagebox.showwarning("Sense resultat", "Google no ha retornat cap resultat usable.")
            return

        record.font = source_url
        if self.current_index is not None:
            self.suspend_dirty = True
            try:
                self.editor_vars["font"].set(source_url)
            finally:
                self.suspend_dirty = False
        self.dirty = False
        self.refresh_tree()
        self.update_validation_message()
        self.set_status("Primera font guardada en la fila actual.")
        self.append_history(f"Font automàtica guardada per a: {record.nom or 'fila seleccionada'}")

    def apply_google_first_result_to_all(self):
        if not self.records:
            messagebox.showwarning("Sense files", "Carrega primer un Excel.")
            return

        updated = 0
        failed = 0
        progress_window, progress_var, progress_status_var, progress_count_var = self.open_progress_window("Buscant fonts automàtiques", len(self.records))
        try:
            for index, record in enumerate(self.records, start=1):
                label = record.nom.strip() or f"fila {index}"
                progress_status_var.set(f"Processant {label}...")
                progress_count_var.set(f"{index - 1} / {len(self.records)}")
                progress_window.update()

                query = self.build_search_query(record)
                if not query:
                    failed += 1
                    progress_var.set(index)
                    progress_count_var.set(f"{index} / {len(self.records)}")
                    progress_window.update()
                    continue
                try:
                    source_url = self.resolve_google_first_result(query)
                except Exception:
                    failed += 1
                    progress_var.set(index)
                    progress_count_var.set(f"{index} / {len(self.records)}")
                    progress_window.update()
                    continue
                if not source_url:
                    failed += 1
                    progress_var.set(index)
                    progress_count_var.set(f"{index} / {len(self.records)}")
                    progress_window.update()
                    continue
                record.font = source_url
                updated += 1
                if self.current_index == index - 1:
                    self.suspend_dirty = True
                    try:
                        self.editor_vars["font"].set(source_url)
                    finally:
                        self.suspend_dirty = False
                    self.dirty = False
                progress_var.set(index)
                progress_count_var.set(f"{index} / {len(self.records)}")
                progress_window.update()
        finally:
            try:
                progress_window.grab_release()
                progress_window.destroy()
            except Exception:
                pass

        self.refresh_tree()
        self.update_validation_message()
        self.set_status(f"Fonts automàtiques aplicades a {updated} files.")
        self.append_history(f"Fonts automàtiques aplicades a {updated} files; sense resultat en {failed}.")
        messagebox.showinfo("Procés completat", f"Fonts aplicades: {updated}\nSense resultat: {failed}")

    def open_source_helper(self):
        if self.current_index is None:
            return
        self.apply_current_record(silent=True)
        record = self.records[self.current_index]
        if record.font.strip():
            webbrowser.open(record.font.strip())
            return
        query = self.build_search_query(record)
        if query:
            webbrowser.open("https://www.google.com/search?q=" + urllib.parse.quote(query))
            return
        fallback = record.localitat.strip() or "ajuntament cultura"
        webbrowser.open(DEFAULT_FACEBOOK_SEARCH + urllib.parse.quote(fallback))

    def open_source(self):
        if self.current_index is None:
            return
        self.apply_current_record(silent=True)
        source = self.records[self.current_index].font.strip()
        if source:
            webbrowser.open(source)
        else:
            messagebox.showinfo("Sense font", "Esta fila encara no té cap font.")

    def ensure_ready_to_submit(self):
        if not self.records:
            messagebox.showwarning("Sense files", "Carrega primer un Excel.")
            return False
        if not self.form_metadata:
            messagebox.showwarning("Sense formulari", "Prem abans en 'Llegir formulari'.")
            return False
        return True

    def invalid_indexes(self, indexes):
        invalid = []
        for idx in indexes:
            errors = self.validate_record(self.records[idx])
            if errors:
                invalid.append((idx, errors))
        return invalid

    def preview_current_payload(self):
        if not self.ensure_ready_to_submit():
            return
        if self.current_index is None:
            messagebox.showinfo("Sense selecció", "Selecciona una fila.")
            return
        self.apply_current_record(silent=True)
        errors = self.validate_record(self.records[self.current_index])
        if errors:
            messagebox.showwarning("Fila incompleta", "\n".join(errors))
            return

        payload = build_payload(self.records[self.current_index], self.form_metadata)
        window = tk.Toplevel(self.root)
        window.title("Previsualització d'enviament")
        window.geometry("900x680")
        window.configure(bg=COLORS["paper"])

        frame = ttk.Frame(window, style="Card.TFrame", padding=16)
        frame.pack(fill="both", expand=True)
        ttk.Label(frame, text="Payload que s'enviarà al formulari", style="Section.TLabel").pack(anchor="w", pady=(0, 8))
        box = tk.Text(frame, wrap="word", font=("Consolas", 10), bg="white", fg=COLORS["charcoal"], relief="flat", padx=12, pady=12)
        box.pack(fill="both", expand=True)
        lines = [f"{key} = {value}" for key, value in payload.items()]
        box.insert("1.0", "\n".join(lines))
        box.configure(state="disabled")

    def submit_selected(self):
        if not self.ensure_ready_to_submit():
            return
        if self.current_index is None:
            return
        self.apply_current_record(silent=True)
        invalid = self.invalid_indexes([self.current_index])
        if invalid:
            messagebox.showwarning("Fila no preparada", "\n".join(invalid[0][1]))
            self.update_validation_summary(self.current_index)
            return
        self.run_submission([self.current_index])

    def submit_all(self):
        if not self.ensure_ready_to_submit():
            return
        self.apply_current_record(silent=True)
        indexes = list(range(len(self.records)))
        invalid = self.invalid_indexes(indexes)
        if invalid:
            lines = []
            for idx, errors in invalid[:8]:
                title = self.records[idx].nom or self.records[idx].localitat or f"fila {idx + 1}"
                lines.append(f"{idx + 1}. {title}: {errors[0]}")
            if len(invalid) > 8:
                lines.append(f"... i {len(invalid) - 8} files més.")
            messagebox.showwarning("Hi ha files per revisar", "\n".join(lines))
            self.refresh_tree()
            return
        self.run_submission(indexes)

    def run_submission(self, indexes):
        if self.worker and self.worker.is_alive():
            messagebox.showinfo("En curs", "Ja hi ha un enviament en marxa.")
            return

        def task():
            ok = 0
            for idx in indexes:
                record = self.records[idx]
                try:
                    payload = build_payload(record, self.form_metadata)
                    body = urllib.parse.urlencode(payload).encode("utf-8")
                    request = urllib.request.Request(
                        self.form_metadata["response_url"],
                        data=body,
                        headers={"User-Agent": "Mozilla/5.0"},
                        method="POST",
                    )
                    with urllib.request.urlopen(request, timeout=30) as response:
                        if response.status != 200:
                            raise ValueError(f"Resposta HTTP {response.status}")
                    record.status = "Enviat"
                    record.status_detail = f"Enviat correctament el {self.now_stamp()}."
                    ok += 1
                    self.root.after(0, lambda idx=idx, name=(record.nom or record.localitat or f"fila {idx + 1}"): self.append_history(f"Enviada correctament: {name}."))
                except Exception as exc:
                    record.status = "Error"
                    record.status_detail = self.friendly_error_message(exc)
                    self.root.after(0, lambda idx=idx, detail=record.status_detail, name=(record.nom or record.localitat or f"fila {idx + 1}"): self.append_history(f"Error en {name}: {detail}"))
                self.root.after(0, self.refresh_tree)
                self.root.after(0, lambda idx=idx: self.reselect_index(idx))
            self.root.after(0, lambda: self.set_status(f"Enviament acabat: {ok}/{len(indexes)} correctes."))

        self.worker = threading.Thread(target=task, daemon=True)
        self.worker.start()
        self.set_status("Enviant formularis...")

    def reselect_index(self, index):
        iid = str(index)
        if iid in self.tree.get_children():
            self.tree.selection_set(iid)
            self.tree.focus(iid)
            self.tree.see(iid)
            self.populate_editor(index)

    def friendly_error_message(self, exc):
        text = str(exc)
        lowered = text.lower()
        if "timed out" in lowered or "timeout" in lowered:
            return "Temps d'espera esgotat en enviar el formulari."
        if "http" in lowered:
            return text
        if "urlopen error" in lowered:
            return "No s'ha pogut connectar amb internet."
        return text

    def open_review_mode(self):
        if self.current_index is None:
            messagebox.showinfo("Sense selecció", "Selecciona una fila abans d'obrir la revisió ampla.")
            return
        self.apply_current_record(silent=True)

        window = tk.Toplevel(self.root)
        window.title("Revisió ampla")
        window.geometry("1100x820")
        window.configure(bg=COLORS["cream"])

        outer = ttk.Frame(window, style="Card.TFrame", padding=18)
        outer.pack(fill="both", expand=True)

        temp_vars = {key: tk.StringVar(value=self.editor_vars[key].get()) for key in FIELD_LABELS}
        temp_texts = {}
        row_index = 0

        for key, label in FIELD_LABELS.items():
            ttk.Label(outer, text=label, style="Body.TLabel").grid(row=row_index, column=0, sticky="nw", pady=6, padx=(0, 10))
            if key in LONG_TEXT_FIELDS:
                box = tk.Text(outer, height=3 if key != "font" else 4, wrap="word", font=("Segoe UI", 10))
                box.insert("1.0", temp_vars[key].get())
                box.grid(row=row_index, column=1, sticky="ew", pady=6)
                temp_texts[key] = box
            elif key in FIELD_OPTIONS:
                ttk.Combobox(outer, textvariable=temp_vars[key], values=FIELD_OPTIONS[key], state="readonly", width=60).grid(row=row_index, column=1, sticky="ew", pady=6)
            else:
                ttk.Entry(outer, textvariable=temp_vars[key], width=70).grid(row=row_index, column=1, sticky="ew", pady=6)
            row_index += 1

        outer.columnconfigure(1, weight=1)

        button_bar = ttk.Frame(outer, style="Card.TFrame")
        button_bar.grid(row=row_index, column=0, columnspan=2, sticky="ew", pady=(16, 0))

        def save_and_close():
            for key, box in temp_texts.items():
                temp_vars[key].set(box.get("1.0", "end-1c").strip())
            self.suspend_dirty = True
            try:
                for key in FIELD_LABELS:
                    self.editor_vars[key].set(temp_vars[key].get())
            finally:
                self.suspend_dirty = False
            self.apply_current_record(silent=True)
            window.destroy()
            self.set_status("Canvis guardats des de la revisió ampla.")

        ttk.Button(button_bar, text="Guardar i tancar", command=save_and_close, style="Neutral.TButton").pack(side="right", padx=4)


def main():
    root = tk.Tk()
    FormFillerApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
